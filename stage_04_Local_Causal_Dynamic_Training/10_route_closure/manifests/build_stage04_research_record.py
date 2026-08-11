"""Create the Chinese Stage 04 Research Record DOCX with fixed design tokens."""

from __future__ import annotations

import json
from pathlib import Path
import sys

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE=Path(__file__).resolve(); CLOSURE=HERE.parents[1]; STAGE04=HERE.parents[2]; ROOT=HERE.parents[3]
OUT=STAGE04/"documents/Stage_04_Research_Record.docx"; ASSET=CLOSURE/"stage04_research_record/stage04_timeline.png"
BLUE=RGBColor(46,116,181); DARK_BLUE=RGBColor(31,77,120); NAVY=RGBColor(32,55,72); MUTED=RGBColor(90,98,108); LIGHT=RGBColor(232,238,245); PALE=RGBColor(244,246,249); RED=RGBColor(155,28,28); GOLD=RGBColor(122,90,0)
FONT="Calibri"; CJK="Heiti SC"; WIDTH_DXA=9360; INDENT_DXA=120


def load(path:Path): return json.loads(path.read_text(encoding="utf-8"))
def set_font(run,size=None,bold=None,color=None,italic=None,name=FONT):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name); run._element.rPr.rFonts.set(qn("w:hAnsi"),name); run._element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if color is not None: run.font.color.rgb=color


def set_repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); trPr.append(el)


def set_cell_margins(cell,top=80,start=120,bottom=80,end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        tag=tcMar.find(qn(f"w:{side}"))
        if tag is None: tag=OxmlElement(f"w:{side}"); tcMar.append(tag)
        tag.set(qn("w:w"),str(val)); tag.set(qn("w:type"),"dxa")


def set_table_geometry(table,widths:list[int]):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr; tblW=tblPr.first_child_found_in("w:tblW"); tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    tblInd=tblPr.first_child_found_in("w:tblInd")
    if tblInd is None: tblInd=OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"),str(INDENT_DXA)); tblInd.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcW=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW"); tcW.set(qn("w:w"),str(widths[i])); tcW.set(qn("w:type"),"dxa"); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell,fill):
    shd=cell._tc.get_or_add_tcPr().first_child_found_in("w:shd")
    if shd is None: shd=OxmlElement("w:shd"); cell._tc.get_or_add_tcPr().append(shd)
    shd.set(qn("w:fill"),fill)


def add_table(doc,caption:str,headers:list[str],rows:list[list[str]],widths:list[int]):
    cp=doc.add_paragraph(style="Caption"); cp.paragraph_format.keep_with_next=True; set_font(cp.add_run(caption),9.5,bold=True,color=DARK_BLUE)
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.text=""; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(h),9,bold=True,color=NAVY); shade(cell,"E8EEF5")
    set_repeat_header(table.rows[0])
    for values in rows:
        cells=table.add_row().cells
        for i,val in enumerate(values):
            cells[i].text=""; p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i==0 or len(str(val))<18 else WD_ALIGN_PARAGRAPH.LEFT; set_font(p.add_run(str(val)),8.5)
    set_table_geometry(table,widths); doc.add_paragraph().paragraph_format.space_after=Pt(2); return table


def add_bullet(doc,text,level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2"); p.paragraph_format.left_indent=Inches(0.375+0.25*level); p.paragraph_format.first_line_indent=Inches(-0.188); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25; set_font(p.add_run(text),11); return p


def add_para(doc,text,bold_lead=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.25
    if bold_lead and text.startswith(bold_lead): set_font(p.add_run(bold_lead),11,bold=True,color=NAVY); set_font(p.add_run(text[len(bold_lead):]),11)
    else: set_font(p.add_run(text),11)
    return p


def add_heading(doc,text,level=1):
    p=doc.add_paragraph(text,style=f"Heading {level}"); p.paragraph_format.keep_with_next=True; return p


def add_field(paragraph,instruction,placeholder=""):
    run=paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin"); instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=instruction; sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate"); txt=OxmlElement("w:t"); txt.text=placeholder; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end"); run._r.extend([begin,instr,sep,txt,end]); return run


def set_alt(picture_run,description):
    for docPr in picture_run._element.xpath(".//wp:docPr"): docPr.set("descr",description); docPr.set("title","Stage 04 evidence timeline")


def add_equation(doc,text,label):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6); r=p.add_run(text); set_font(r,11,name="Cambria Math"); set_font(p.add_run(f"    ({label})"),10,color=MUTED)


def make_timeline():
    ASSET.parent.mkdir(parents=True,exist_ok=True)
    stages=["04A","04A-V","04B","04C","04C-R","04C-S"]; labels=["合同","验证","参考池","梯度失败","混合归因","路线封闭"]; colors=["#2E74B5","#2E74B5","#2E74B5","#9B1C1C","#7A5A00","#1F4D78"]
    image=Image.new("RGB",(1800,400),"white"); draw=ImageDraw.Draw(image); xs=[150+i*300 for i in range(6)]; y=150
    draw.line((xs[0],y,xs[-1],y),fill="#AAB4C0",width=8)
    latin=ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf",34); cjk=ImageFont.truetype("/System/Library/Fonts/STHeiti Medium.ttc",30)
    for x,stage,label,color in zip(xs,stages,labels,colors):
        draw.ellipse((x-72,y-72,x+72,y+72),fill=color,outline="white",width=6)
        box=draw.textbbox((0,0),stage,font=latin); draw.text((x-(box[2]-box[0])/2,y-(box[3]-box[1])/2-4),stage,font=latin,fill="white")
        box=draw.textbbox((0,0),label,font=cjk); draw.text((x-(box[2]-box[0])/2,y+104),label,font=cjk,fill="#263442")
    image.save(ASSET,dpi=(180,180))


def configure(doc):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(0.492)
    styles=doc.styles
    n=styles["Normal"]; n.font.name=FONT; n.font.size=Pt(11); n._element.rPr.rFonts.set(qn("w:eastAsia"),CJK); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.25
    tokens={1:(16,BLUE,18,10),2:(13,BLUE,14,7),3:(12,DARK_BLUE,10,5)}
    for level,(size,color,before,after) in tokens.items():
        s=styles[f"Heading {level}"]; s.font.name=FONT; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color; s._element.rPr.rFonts.set(qn("w:eastAsia"),CJK); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    cap=styles["Caption"]; cap.font.name=FONT; cap.font.size=Pt(9.5); cap.font.color.rgb=DARK_BLUE; cap._element.rPr.rFonts.set(qn("w:eastAsia"),CJK); cap.paragraph_format.space_before=Pt(4); cap.paragraph_format.space_after=Pt(4)
    for name in ("List Bullet","List Bullet 2"):
        s=styles[name]; s.font.name=FONT; s.font.size=Pt(11); s._element.rPr.rFonts.set(qn("w:eastAsia"),CJK); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; header.paragraph_format.space_after=Pt(0); set_font(header.add_run("SPH-PIO-PoC  |  Stage 04 Research Record"),8.5,color=MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(footer.add_run("Stage 04 路线封闭研究记录  ·  "),8.5,color=MUTED); add_field(footer,"PAGE","1"); set_font(footer.add_run(" / "),8.5,color=MUTED); add_field(footer,"NUMPAGES","—")
    doc.core_properties.title="Stage 04 Research Record"; doc.core_properties.subject="Local-causal training route closure and evidence record"; doc.core_properties.author="SPH-PIO-PoC Project"; doc.core_properties.keywords="SPH, task gradient, RK2, verification, negative evidence"


def main():
    ledger=load(CLOSURE/"status_ledger/stage04_status_ledger.json"); evidence=load(CLOSURE/"evidence_matrix/stage04_evidence_matrix.json"); boundary=load(CLOSURE/"failure_boundary/stage04_task_signal_failure_boundary.json"); innovation=load(CLOSURE/"innovation_register/stage04_innovation_register.json"); claims=load(CLOSURE/"claim_boundary/stage04_claim_boundary.json"); publication=load(CLOSURE/"publication_delta/publication_option_update.json"); freeze=load(STAGE04/"09_manifests/stage04cs_input_freeze_manifest.json")
    make_timeline(); doc=Document(); configure(doc)
    # Editorial-cover override atop compact_reference_guide.
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(92); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run("研究记录 · 路线封闭"),11,bold=True,color=GOLD)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(8); set_font(p.add_run("Stage 04 Research Record"),30,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4); set_font(p.add_run("局部因果动态训练路线：从假设、参考池到任务信号边界"),15,bold=True,color=DARK_BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(44); set_font(p.add_run("SPH-PIO-PoC · Stage 04C-S · 2026-08-06"),10.5,color=MUTED)
    pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=pic.add_run(); run.add_picture(str(ASSET),width=Inches(6.15)); set_alt(run,"Stage 04A 到 Stage 04C-S 的证据时间线：合同、验证、参考池、梯度失败、混合归因与路线封闭。")
    cp=doc.add_paragraph(style="Caption"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(cp.add_run("图 1  Stage 04 证据链与非覆盖式状态演进"),9.5,bold=True,color=DARK_BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(26); set_font(p.add_run("最终边界：STAGE04_ROUTE_PAUSED_TASK_SIGNAL_BOUNDARY_COMPLETE"),10.5,bold=True,color=RED)
    doc.add_page_break()
    add_heading(doc,"内容索引",1)
    toc_items=[
        "执行摘要", "1. Stage 04 新假设", "2. Stage 03 边界与继承关系",
        "3. K=1 local-causal 训练定义", "4. 新 reference formula pool",
        "5. Analytic、DOP853 与 topology 资格", "6. Lineage split 与 sealed test",
        "7. Task-aligned component-vector loss", "8. D1/D2/D3 parameter groups",
        "9. 864 个正式 probes", "10. Reverse、JVP 与 central FD",
        "11. All-near-zero 资格结果", "12. Full parameter gradients",
        "13. Exact residual–Jacobian factorization", "14. Random-direction projection",
        "15. Network sensitivity chain", "16. RK2 time attenuation",
        "17. Mixed / unresolved attribution", "18. 为什么没有训练",
        "19. Supported / conditional / unsupported claims",
        "20. Stage 00–04 publication implications", "附录 A. Evidence matrix 摘要",
        "附录 B. 潜在创新登记", "附录 C. Artifact / hash index",
    ]
    toc_rows=[]
    half=(len(toc_items)+1)//2
    for i in range(half):
        toc_rows.append([toc_items[i], toc_items[i+half] if i+half < len(toc_items) else ""])
    add_table(doc,"表 0  研究记录内容索引",["前半部分","后半部分"],toc_rows,[4680,4680])
    add_para(doc,"阅读说明：本记录保留每个阶段的独立 verdict。后续归因只扩展证据，不覆盖早期失败；所有“潜在创新”均需独立文献核验。")
    doc.add_page_break()

    add_heading(doc,"执行摘要",1)
    add_para(doc,"Stage 04 建立了新的 K=1 local-causal 动态训练假设、10 个 formula lineages 的角色预分配 reference pool，以及面向真实 optimizer 参数组的 task-aligned 梯度资格体系。Reference pool 完整通过，但 864 个正式 probes 均因三个 loss 分量同时低于可辨识门而失败。")
    add_para(doc,"Stage 04C-R 进一步证明网络输出、参数全梯度与状态 Jacobian 并非系统性死亡；失败由 MSE residual scale、随机组方向投影和合约内 RK2 时间尺度共同影响，且仍有 604 个 component rows 未能唯一解析。故 Stage 04D=false，训练未授权、未执行。")
    add_table(doc,"表 1  Stage 04 状态总览",["阶段","精确状态","下游边界"],[[r["stage"],r["exact_status"],r["downstream_authorization"]] for r in ledger["rows"]],[1300,3600,4460])

    add_heading(doc,"1. Stage 04 新假设",1)
    add_para(doc,"核心假设不是直接进行长时训练，而是先判断一个完整 K=1 RK2 state transition 的 task-aligned loss 是否能在实际 optimizer 参数空间中产生可检测、可复现并与有限差分一致的梯度。")
    add_bullet(doc,"局部因果：仅使用当前状态、已接受 prehistory 与当前/中点 graph。")
    add_bullet(doc,"任务对齐：梯度对象直接来自下一步位置、速度、密度状态误差。")
    add_bullet(doc,"资格优先：optimizer、training、checkpoint selection 与性能评价均置于梯度资格之后。")

    add_heading(doc,"2. Stage 03 边界与继承关系",1)
    add_para(doc,"Stage 03D 的多步 AD/FD 与 history 梯度未资格化，Stage 03D-R 给出 mixed/unresolved 归因，Stage 03D-S 将路线暂停在梯度边界。Stage 04 不修复或覆盖这些结论，而是通过新的 TRAIN formula families 与 K=1 optimizer-variable loss 重新建立独立假设。")
    add_para(doc,"不变边界：Stage 03E 未授权；Stage 04 不得把 K=1 的实现一致性写成多步 rollout 梯度已经解决。")

    add_heading(doc,"3. K=1 local-causal 训练定义",1)
    add_para(doc,"每个任务样本执行 start → midpoint → accept 的完整显式中点 RK2。Start 与 midpoint 分别重建 reciprocal graph；midpoint token 为 ephemeral；D2/D3 仅在 accept 后提交一次 history，midpoint commit count=0。")
    add_equation(doc,"Sⁿ → k₁(Sⁿ, Hⁿ; θ) → Sⁿ⁺¹ᐟ² → k₂(Sⁿ⁺¹ᐟ², Hⁿ; θ) → Sⁿ⁺¹","3.1")

    add_heading(doc,"4. 新 reference formula pool",1)
    add_para(doc,"Stage 04B 构建 10 个公式血统，覆盖轴向压缩、旋转单元、斜向纵/横模态、多模态交叉与各向异性双模态。每个 lineage 具有 LOW/MAIN 变体与 N8/N12/N16 exact trajectories。")
    add_bullet(doc,"Analytic qualification：20/20 PASS。")
    add_bullet(doc,"Exact trajectories：60/60 PASS。")
    add_bullet(doc,"DOP853：20/20 PASS；fixed topology：10/10。")

    add_heading(doc,"5. Analytic、DOP853 与 topology 资格",1)
    add_para(doc,"解析公式、半离散参考与图拓扑分别审计，避免将 analytic identity、时间积分一致性和固定拓扑平滑性混写为一个门。Stage 04C 的所有扰动路径再次验证了 start/midpoint/accepted edge identity，topology-changing epsilon=0。")

    add_heading(doc,"6. Lineage split 与 sealed test",1)
    add_para(doc,"角色在科学结果之前冻结为 6 TRAIN、2 VALIDATION、2 SEALED_TEST。Stage 04C 与 04C-R 仅使用 TRAIN；validation target 及 sealed formula/state/target/origin 均保持零解码。")
    add_table(doc,"表 2  角色与访问边界",["角色","Lineages","Stage 04C/CR 访问"],[["TRAIN","LCDF_01, 04, 05, 06, 07, 08","允许"],["VALIDATION","LCDF_02, 09","禁止解码"],["SEALED_TEST","LCDF_03, 10","公式/状态/目标/来源/图序列均禁止"]],[1500,3300,4560])

    add_heading(doc,"7. Task-aligned component-vector loss",1)
    add_para(doc,"正式梯度对象是三个 loss components，而不是事后选择的加权和。Stage 04D 如未来提出固定非负权重，必须在新合同中进行线性组合核验。")
    add_equation(doc,"Lₓ = meanᵢ ‖minimum_image(xᵢⁿ⁺¹ − xᵢ,refⁿ⁺¹)‖² / L²","7.1")
    add_equation(doc,"Lᵥ = meanᵢ ‖vᵢⁿ⁺¹ − vᵢ,refⁿ⁺¹‖² / cₛ²","7.2")
    add_equation(doc,"Lρ = meanᵢ (ρᵢⁿ⁺¹ − ρᵢ,refⁿ⁺¹)² / ρ₀²","7.3")

    add_heading(doc,"8. D1/D2/D3 parameter groups",1)
    add_para(doc,"参数组从 Stage 03C 实际 tensor paths 解析，所有 trainable elements 恰好分配一次。D3 合并的 in_proj_weight/bias 按 Q/K/V 不重叠 slices 冻结。")
    add_table(doc,"表 3  参数组与规模",["Arm","Groups","参数总数"],[["D1","TOKEN_ENCODER; PAIR_HEAD","5,762"],["D2","TOKEN_ENCODER; GRU; PAIR_HEAD","12,098"],["D3","TOKEN_ENCODER; ATTENTION Q/K/V/O; FEED_FORWARD; PAIR_HEAD","22,978"]],[1200,5960,2200])

    add_heading(doc,"9. 864 个正式 probes",1)
    add_para(doc,"每个 arm 使用 6 lineages × 2 variants × 2 preregistered origins × 3 model seeds = 72 contexts，再与实际参数组组合形成 D1 144、D2 216、D3 504，总计 864 probes。所有 origin 与方向在 state decode 前通过 SHA-256 规则冻结。")

    add_heading(doc,"10. Reverse、JVP 与 central FD",1)
    add_para(doc,"每个 probe 对 Lx/Lv/Lrho 分别执行 reverse VJP、真实 JVP 与 5 点 central-FD ladder；所有路径重复两次并处于 CPU float64 / explicit math-SDPA context。")
    add_bullet(doc,"Reverse/JVP component comparisons：2592/2592 PASS。")
    add_bullet(doc,"FD plus/minus paths：17280；topology changes：0。")
    add_bullet(doc,"Parameter mutation、access violation、finite/safety failure：0。")

    add_heading(doc,"11. All-near-zero 资格结果",1)
    add_para(doc,"尽管 AD/JVP/FD 实现一致，2592 个 component directional derivatives 全部低于合同分辨率；864 probes 均为 all-near-zero，因此没有任何 probe 能满足“至少一个非零 stable component”。D1/D2/D3 pass rate 均为 0%，qualified parameter groups=0。")
    add_para(doc,"允许结论：预注册 K=1 task-aligned gradient qualification 未能在所有必要参数组上建立足够可检测的非零 task-loss sensitivities。禁止结论：模型或 Transformer 无法训练。")

    add_heading(doc,"12. Full parameter gradients",1)
    add_para(doc,"Stage 04C-R 对 216 formal contexts 计算全部参数梯度，并按 864 group rows 报告 L2/RMS/Linf、符号平衡与 decade histogram。重复计算完全确定，parameter hashes 前后 bitwise 相等。")
    add_bullet(doc,"位置 full gradients 通常低于 1e−14。")
    add_bullet(doc,"速度 full gradients 通常可检测，部分达到或超过 1e−12。")
    add_bullet(doc,"密度 full gradients 位于检测边界，跨组表现混合。")

    add_heading(doc,"13. Exact residual–Jacobian factorization",1)
    add_para(doc,"对原冻结方向 d，定义无量纲 residual e 与 state JVP z=Jd。所有 component 均满足以下恒等式，并与历史 reverse derivative 精确重建。")
    add_equation(doc,"dL_c / dε = 2 · mean(e_c · z_c)","13.1")
    add_para(doc,"2592/2592 factorization PASS，最大绝对重建误差处于 1e−24 量级。这排除了 Autograd 矛盾，并将极小任务梯度分解为 residual、state sensitivity 与 alignment 三因素。")

    add_heading(doc,"14. Random-direction projection",1)
    add_para(doc,"当 full group gradient 明显非零时，原单位 Rademacher 方向的投影会按参数维数约以 1/√D 稀释。672/2592 components 被归为 projection-primary；detectable gradients 的 median scaled projection=0.655，与随机方向理论一致，但只解释 25.9%，不足以形成单一总体归因。")

    add_heading(doc,"15. Network sensitivity chain",1)
    add_para(doc,"同一参数方向上的 hidden → alpha/beta → pair force → nodal correction acceleration → midpoint state → accepted state → final loss 链全部 finite。Hidden、coefficient、force 与 acceleration JVP 均显著非零；tanh saturation fraction=0；final head 非零；无 hidden collapse。")
    add_para(doc,"因此 dead network、zero head、saturation dead zone 与 parameterization-wide dead sensitivity 均被正式排除。")

    add_heading(doc,"16. RK2 time attenuation",1)
    add_para(doc,"网络 acceleration JVP 进入 accepted state 时按积分合同被时间尺度压缩。中位比值 V_accept/(Δt A_mid)≈1，X_accept/(Δt² A_mid)≈0.5；这是显式中点 RK2 的预期关系，不是实现缺陷。")
    add_equation(doc,"δvⁿ⁺¹ ≈ Δt · δa_mid,     δxⁿ⁺¹ ≈ ½Δt² · δa_start","16.1")

    add_heading(doc,"17. Mixed / unresolved attribution",1)
    add_table(doc,"表 4  归因构成",["Primary reason","Rows","Share"],[["TASK_RESIDUAL_TOO_SMALL","1316","50.8%"],["GROUP_DIRECTION_PROJECTION_DILUTION","672","25.9%"],["UNRESOLVED","604","23.3%"]],[4700,1800,2860])
    add_para(doc,"位置 rows 全部 residual-limited；速度 rows 以 projection dilution 为主；密度 rows 在 residual-small 与 unresolved 之间分裂。无单因达到预注册的 80% 总体门，因此唯一总体状态为 TASK_GRADIENT_FAILURE_MIXED_OR_UNRESOLVED。")

    add_heading(doc,"18. 为什么没有训练",1)
    add_para(doc,"Stage 04D 的必要授权是 Stage 04C task-gradient qualification PASS。该条件未满足，且 Stage 04C-R 没有形成唯一可前瞻修正的归因分支。因此直接缩放 loss、放宽 1e−10 门、改用 full-gradient training 或只训练“表现较好”的 arm 都属于事后改合同。")
    add_bullet(doc,"Optimizer instances=0；optimizer steps=0；parameter updates=0。")
    add_bullet(doc,"Training runs=0；neural rollout=0；performance evaluations=0。")
    add_bullet(doc,"Stage 04D authorization=false；Training=NOT_AUTHORIZED / NOT_EXECUTED。")

    add_heading(doc,"19. Supported / conditional / unsupported claims",1)
    claim_rows=[]
    for level in ("SUPPORTED","CONDITIONAL","UNSUPPORTED"):
        for claim in claims[level]: claim_rows.append([level,claim])
    add_table(doc,"表 5  Claim boundary",["等级","声明"],claim_rows,[2100,7260])

    add_heading(doc,"20. Stage 00–04 publication implications",1)
    add_para(doc,"本阶段仅更新证据，不作投稿合并/拆分的最终决定。")
    add_table(doc,"表 6  初步投稿选项",["Option","Current support","CMAME readiness","边界"],[[o["option"],o["current_support"],o["CMAME_readiness"],o["reason"]] for o in publication["options"]],[900,2100,2800,3560])
    add_para(doc,"Option C 当前具有最强的方法论主线一致性，但仍需 literature verification 与跨案例普适性论证；不能因为叙事更完整就将其写成最终投稿选择。")

    add_heading(doc,"附录 A. Evidence matrix 摘要",1)
    add_table(doc,"表 A1  Stage 04 证据状态",["Category","Evidence","Status"],[[r["category"],r["evidence"],r["status"]] for r in evidence["rows"]],[2300,5260,1800])

    add_heading(doc,"附录 B. 潜在创新登记",1)
    for row in innovation["rows"]: add_bullet(doc,f"{row['id']}. {row['innovation']} — POTENTIAL_NOVELTY_REQUIRES_LITERATURE_VERIFICATION")

    add_heading(doc,"附录 C. Artifact / hash index",1)
    artifacts=[
        ("Stage04A final","stage_04_Local_Causal_Dynamic_Training/09_manifests/stage04a_final_manifest.json"),("Stage04A verification","stage_04_Local_Causal_Dynamic_Training/00_stage04a_verification/manifests/stage04a_target_verification_manifest.json"),("Stage04B final","stage_04_Local_Causal_Dynamic_Training/09_manifests/stage04b_final_manifest.json"),("Stage04C final","stage_04_Local_Causal_Dynamic_Training/09_manifests/stage04c_final_manifest.json"),("Stage04C-R final","stage_04_Local_Causal_Dynamic_Training/09_manifests/stage04cr_final_manifest.json"),("Stage04C 864 matrix","stage_04_Local_Causal_Dynamic_Training/05_task_aligned_gradient/stage04c/results/formal_864_probe_results.json"),("Stage04C-R factorization","stage_04_Local_Causal_Dynamic_Training/05_task_aligned_gradient/stage04cr/loss_factorization/exact_loss_factorization.json"),("Stage04 status ledger","stage_04_Local_Causal_Dynamic_Training/10_route_closure/status_ledger/stage04_status_ledger.json"),("Stage04 evidence matrix","stage_04_Local_Causal_Dynamic_Training/10_route_closure/evidence_matrix/stage04_evidence_matrix.json"),("Project delta manifest","project_wide_synthesis/11_stage04_update_interface/stage04_completed_delta/stage04_delta_manifest.json")]
    rows=[]
    import hashlib
    for label,path in artifacts:
        p=ROOT/path; rows.append([label,path,"sha256:"+hashlib.sha256(p.read_bytes()).hexdigest()])
    add_table(doc,"表 C1  关键 artifact identities",["Artifact","Path","SHA-256"],rows,[1900,4300,3160])
    add_para(doc,f"完整历史冻结包含 {freeze['historical_file_count']} 个可读文件；另有 {freeze['protected_private_file_count']} 个 validation/sealed 私有文件保持不可读，其身份由 Stage 04B seal/trajectory/role manifests 锚定，受保护 payload read count=0。")

    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT)
    tokens={"preset":"compact_reference_guide","header_override":"editorial_cover","page":{"size":"Letter","margins_in":1.0,"header_footer_in":0.492,"usable_width_dxa":9360},"body":{"font":"Calibri","eastAsia":"Heiti SC","size_pt":11,"after_pt":6,"line_spacing":1.25},"headings":{"h1":[16,"#2E74B5",18,10],"h2":[13,"#2E74B5",14,7],"h3":[12,"#1F4D78",10,5]},"lists":{"marker_in":0.187,"text_indent_in":0.375,"hanging_in":0.188,"after_pt":4,"line_spacing":1.25},"tables":{"width_dxa":9360,"indent_dxa":120,"margins_dxa":[80,120,80,120],"header_fill":"#E8EEF5"},"named_overrides":{"cover_title":{"size_pt":30,"color":"#203748"},"dense_artifact_table":{"font_pt":8.5}}}
    (CLOSURE/"stage04_research_record/design_tokens.json").write_text(json.dumps(tokens,indent=2,ensure_ascii=False)+"\n",encoding="utf-8")
    print(json.dumps({"docx":str(OUT.relative_to(ROOT)),"sections":20,"appendices":3,"figure":str(ASSET.relative_to(ROOT))}))


if __name__=="__main__": main()
