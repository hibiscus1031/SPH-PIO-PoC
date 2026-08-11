#!/usr/bin/env python3
"""Generate the Chinese Stage 02 Research Record from machine manifests."""

from __future__ import annotations

import json
import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[3]
STAGE = REPO / "stage_02_Particle_Interaction_Operator"
ROOT = STAGE / "08_route_closure"
OUT = STAGE / "documents/Stage_02_Research_Record.docx"
ASSETS = ROOT / "figure_plan/record_assets"
SKILL = Path("/Users/xiejinbo/.codex/plugins/cache/openai-primary-runtime/documents/26.802.11031/skills/documents")
sys.path.insert(0, str(SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402

ledger = json.loads((ROOT / "status_ledger/stage02_complete_status_ledger.json").read_text())
matrix = json.loads((ROOT / "evidence_matrix/stage02_complete_evidence_matrix.json").read_text())
taxonomy = json.loads((ROOT / "failure_taxonomy/stage02_failure_taxonomy.json").read_text())
claims = json.loads((ROOT / "claim_boundary/stage02_claim_boundary.json").read_text())
assessment = json.loads((ROOT / "manuscript_assessment/stage02_manuscript_readiness.json").read_text())
figure_plan = json.loads((ROOT / "figure_plan/stage02_figure_and_table_plan.json").read_text())
future = json.loads((ROOT / "future_branches/stage03_branch_decision_design.json").read_text())
freeze = json.loads((ROOT / "freeze/stage02ms_historical_freeze_manifest.json").read_text())

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183B56"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "2F6B4F"


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Heiti SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT, size: float = 8.7) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]], weights: list[float], *, font_size: float = 8.7) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
        shade(table.rows[0].cells[i], LIGHT_BLUE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or (isinstance(value, (int, float)) and not isinstance(value, bool)) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align, size=font_size)
            if row_index % 2:
                shade(cells[i], "FAFBFC")
    apply_table_geometry(table, column_widths_from_weights(weights, 9360), table_width_dxa=9360, indent_dxa=120, cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120})
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return table


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=9, color=MID_GRAY, italic=True)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))


def add_equation(doc: Document, expression: str, label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(expression + (f"    ({label})" if label else ""))
    set_font(r, size=11.5, italic=True, name="Cambria Math")


def add_callout(doc: Document, label: str, text: str, color: str = NAVY) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + "  ")
    set_font(r, size=10, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, size=10)


def add_picture_with_alt(doc: Document, path: Path, alt_text: str, width: float = 6.35) -> None:
    shape = doc.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_static_toc(doc: Document, entries: list[tuple[str, str]]) -> None:
    """Add a renderer-independent one-page table of contents."""
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "8640")
        tabs.append(tab)
        p._p.get_or_add_pPr().append(tabs)
        r = p.add_run(f"{title}\t{page}")
        set_font(r, size=8.8, color=NAVY)


def add_page_field(paragraph) -> None:
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.extend([begin, instr, separate, t, end])
    set_font(r, size=9, color=MID_GRAY)


def font_path() -> str:
    candidates = ["/System/Library/Fonts/PingFang.ttc", "/System/Library/Fonts/STHeiti Light.ttc", "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"]
    return next(path for path in candidates if Path(path).is_file())


FONT_PATH = font_path()


def flow_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 720), "white")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(FONT_PATH, 46, index=0)
    body = ImageFont.truetype(FONT_PATH, 27, index=0)
    small = ImageFont.truetype(FONT_PATH, 23, index=0)
    draw.text((70, 35), "V&V-first PIO 资格化流程与终止边界", font=title, fill="#183B56")
    labels = [
        ("Reference", "Fourier / analytic\n独立一致"),
        ("Target", "空间归因\n守恒 scope"),
        ("Dataset", "blind lineage\n10/5/5 split"),
        ("Architecture", "K1/K2\n结构门 PASS"),
        ("Protocol v0.1", "9 runs\nNOT QUALIFIED"),
        ("Attribution", "conditioning\ndiagnostic"),
        ("Protocol v0.2", "9 runs\nNOT QUALIFIED"),
        ("Closure", "route terminated\nno Stage 02N"),
    ]
    x0, y0, width, height, gap = 45, 215, 185, 260, 34
    for i, (head, detail) in enumerate(labels):
        x = x0 + i * (width + gap)
        fill = "#E8EEF5" if i < 4 else ("#FFF4F4" if i in (4, 6, 7) else "#FFF8E8")
        outline = "#2E74B5" if i < 4 else ("#9B1C1C" if i in (4, 6, 7) else "#7A5A00")
        draw.rounded_rectangle((x, y0, x + width, y0 + height), radius=18, fill=fill, outline=outline, width=4)
        hb = draw.textbbox((0, 0), head, font=body)
        draw.text((x + (width - (hb[2] - hb[0])) / 2, y0 + 30), head, font=body, fill=outline)
        for j, line in enumerate(detail.split("\n")):
            box = draw.textbbox((0, 0), line, font=small)
            draw.text((x + (width - (box[2] - box[0])) / 2, y0 + 112 + j * 42), line, font=small, fill="#263238")
        if i < len(labels) - 1:
            ax = x + width + 4
            ay = y0 + height / 2
            draw.line((ax, ay, ax + gap - 8, ay), fill="#6B7280", width=5)
            draw.polygon([(ax + gap - 8, ay), (ax + gap - 24, ay - 10), (ax + gap - 24, ay + 10)], fill="#6B7280")
    draw.text((60, 600), "硬边界：Stage 01 V2 FAIL 保留；validation/test gate PASS 不覆盖 train-fit failure；rollout 与 solver-in-loop 未执行。", font=small, fill="#9B1C1C")
    image.save(path)


def timeline_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(FONT_PATH, 46, index=0)
    body = ImageFont.truetype(FONT_PATH, 25, index=0)
    small = ImageFont.truetype(FONT_PATH, 20, index=0)
    draw.text((70, 35), "Stage 02 状态链：后续状态不覆盖历史失败", font=title, fill="#183B56")
    milestones = [(row["stage"].replace("Stage ", ""), row["unique_status"]) for row in ledger["rows"]]
    cols = 6
    box_w, box_h, x_gap, y_gap = 245, 155, 40, 60
    for i, (stage, status) in enumerate(milestones):
        row, col = divmod(i, cols)
        x, y = 60 + col * (box_w + x_gap), 150 + row * (box_h + y_gap)
        negative = any(token in status for token in ("NOT_READY", "NOT_QUALIFIED", "FAIL", "TERMINATED"))
        fill = "#FFF4F4" if negative else "#E8EEF5"
        outline = "#9B1C1C" if negative else "#2E74B5"
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=14, fill=fill, outline=outline, width=3)
        draw.text((x + 14, y + 14), stage, font=body, fill=outline)
        words = status.split("_")
        lines = []
        current = ""
        for word in words:
            candidate = word if not current else current + "_" + word
            if draw.textbbox((0, 0), candidate, font=small)[2] > box_w - 28 and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        for j, line in enumerate(lines[:4]):
            draw.text((x + 14, y + 56 + j * 24), line, font=small, fill="#263238")
    image.save(path)


def boundary_figure(path: Path) -> None:
    image = Image.new("RGB", (1800, 850), "white")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(FONT_PATH, 46, index=0)
    head = ImageFont.truetype(FONT_PATH, 31, index=0)
    body = ImageFont.truetype(FONT_PATH, 23, index=0)
    draw.text((70, 35), "证据—主张边界", font=title, fill="#183B56")
    columns = [
        ("SUPPORTED", "#EAF5EF", "#2F6B4F", ["reference-qualified dataset", "pair-force momentum contract", "K1/K2 structural gates", "two static protocols failed"]),
        ("CONDITIONAL", "#FFF8E8", "#7A5A00", ["conditioning contributed", "v0.2 conditioning improved", "local mapping may be difficult"]),
        ("UNSUPPORTED", "#FFF4F4", "#9B1C1C", ["SPH improved / V2 restored", "attention superior / Transformer necessary", "rollout stable / solver accelerated", "arbitrary-flow generalization"]),
    ]
    for i, (heading, fill, outline, items) in enumerate(columns):
        x = 70 + i * 575
        draw.rounded_rectangle((x, 150, x + 510, 760), radius=20, fill=fill, outline=outline, width=4)
        hb = draw.textbbox((0, 0), heading, font=head)
        draw.text((x + (510 - (hb[2] - hb[0])) / 2, 185), heading, font=head, fill=outline)
        y = 285
        for item in items:
            draw.ellipse((x + 35, y + 7, x + 49, y + 21), fill=outline)
            draw.text((x + 68, y), item, font=body, fill="#263238")
            y += 92
    image.save(path)


ASSETS.mkdir(parents=True, exist_ok=True)
flow_path = ASSETS / "record_figure_01_pipeline.png"
timeline_path = ASSETS / "record_figure_02_timeline.png"
boundary_path = ASSETS / "record_figure_03_claim_boundary.png"
flow_figure(flow_path)
timeline_figure(timeline_path)
boundary_figure(boundary_path)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for name, size, color, before, after in (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, DARK_BLUE, 10, 5)):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[list_name].font.name = "Calibri"
    styles[list_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    styles[list_name].font.size = Pt(11)
    styles[list_name].paragraph_format.space_after = Pt(4)
    styles[list_name].paragraph_format.line_spacing = 1.25
styles["Caption"].font.name = "Calibri"
styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
styles["Caption"].font.size = Pt(9)
styles["Caption"].font.italic = True
styles["Caption"].font.color.rgb = RGBColor.from_string(MID_GRAY)

settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
set_font(hp.add_run("SPH-PIO-PoC  |  Stage 02 Research Record"), size=9, color=MID_GRAY)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Stage 02 研究记录  |  "), size=9, color=MID_GRAY)
add_page_field(fp)

# Cover: editorial_cover pattern.
doc.add_paragraph().paragraph_format.space_after = Pt(90)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("SPH-PIO-PoC"), size=12, bold=True, color=GOLD)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_font(p.add_run("Stage 02 研究记录"), size=30, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("粒子相互作用算子：从 reference construction 到静态学习路线证伪"), size=15, color=DARK_BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(44)
set_font(p.add_run("Static Learning Route Closure, Evidence Synthesis and Publication Boundary Assessment"), size=10.5, italic=True, color=MID_GRAY)
add_callout(doc, "最终状态", "STATIC_PAIR_FORCE_FITTING_V02_NOT_QUALIFIED；static PIO learning route terminated；Stage 02N / v0.3 / rollout / solver-in-the-loop 均不授权。", RED)
doc.add_paragraph().paragraph_format.space_after = Pt(55)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("历史文件冻结：1,788  |  Stage 02 状态：22  |  新训练：0  |  新测试：0"), size=10.5, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("生成日期：2026-08-04  |  版本：Stage 02M-S closure record 1.0"), size=9.5, color=MID_GRAY)
page_break(doc)

doc.add_heading("目录", level=1)
add_static_toc(doc, [
    ("摘要", "3"),
    ("1. Stage 02 研究目标与初始假设", "4"),
    ("2. PIO 理论合同", "5"),
    ("3. Target 与 reference qualification", "6"),
    ("4. Conservation-compatible scope", "7"),
    ("5. Dataset construction and leakage control", "8"),
    ("6. Regularity-hard-gate route and termination", "8"),
    ("7. Blind multifamily dataset", "9"),
    ("8. Pair-force architecture", "11"),
    ("9. Static fitting protocol v0.1", "11"),
    ("10. Failure attribution", "11"),
    ("11. Static fitting protocol v0.2", "12"),
    ("12. Final route termination", "12"),
    ("13. Scientific conclusions", "13"),
    ("14. Unsupported claims", "14"),
    ("15. Future research branches", "14"),
    ("16. Artifact and hash index", "15"),
    ("附录 A：完整 Stage 02 状态账本", "18"),
    ("附录 B：证据矩阵", "20"),
    ("附录 C：失败分类与状态语义", "22"),
    ("附录 D：论文可行性与图表包", "23"),
    ("附录 E：Stage 02M-Q checkpoint index", "24"),
    ("记录闭包声明", "24"),
])
page_break(doc)

doc.add_heading("摘要", level=1)
add_body(doc, "Stage 02 研究路线以 Stage 01 的可审计 SPH baseline 和 V2_QUALIFICATION_FAIL 为不变边界，逐步建立 PIO 理论合同、reference hierarchy、空间 target 归因、守恒兼容 scope、blind multifamily dataset、pair-force architecture、两轮预注册静态拟合及其失败归因。本记录从机器 manifest 生成，不删除任何失败阶段，也不把后续 PASS 写成对历史失败的覆盖。")
add_body(doc, "最终证据显示：Fourier 与 analytic reference 在冻结 periodic-vortex scope 内独立一致；最终 blind dataset 具有四个 lineage-disconnected family components、10/5/5 预冻结 split 和 train-only normalization；K1/K2 在 pair antisymmetry、线性动量、O(2) 等变、周期性、可微性和 edge-local resource scaling 合同下 PASS。然而，结构正确性并未带来静态可学性资格。v0.1 与 v0.2 两个九运行协议均未满足冻结 train-fit 门；v0.2 中 K1 为 0/3、K2 为 1/3。")
add_body(doc, "因此，唯一结论是 STATIC_PAIR_FORCE_FITTING_V02_NOT_QUALIFIED，当前 static PIO learning route 终止。rollout 与 solver-in-the-loop 从未授权、从未执行，不能写成动态失败。当前最可行论文主线是 V&V-first methodology 与 negative result 的组合，而不是 Transformer/Attention-corrected solver success。")
add_table(doc, ["记录维度", "闭包结果"], [["历史文件", f"{freeze['historical_file_count']} files, hashes frozen"], ["Stage 02 状态", "22/22 unique, superseded=false"], ["正式静态训练", "v0.1: 9 runs; v0.2: 9 runs"], ["新 Stage 02M-S 训练/测试", "0 / 0"], ["最终路线", "terminated; Stage 02N not authorized"]], [2.0, 4.5], font_size=9.2)

page_break(doc)
doc.add_heading("1. Stage 02 研究目标与初始假设", level=1)
add_body(doc, "Stage 02 的研究问题不是直接证明机器学习可改善 SPH，而是判断：在 reference、target、守恒、数据 lineage、架构对称性、训练协议和测试封存都可审计的条件下，是否存在可资格化的 conservative learned correction operator。")
add_body(doc, "初始工作假设是增量式修正可以保留 baseline 的物理解释和 zero fallback，同时把可归因的空间离散残差映射为粒子或 pair level correction。该假设被拆分为四层可证伪门：reference/target 资格、dataset资格、architecture资格、learning资格。")
add_equation(doc, "Δa(x,t) = a_ref(x,t) − a_SPH(x,t)", "1")
for item in ["Stage 01 结论作为不可修改的外部边界；", "每个后续 PASS 只对本阶段合同有效；", "失败候选必须保留 reason code 与 hash；", "训练成功必须由预注册 train/validation/test 与结构门共同定义；", "未授权的动态证据不能被推断。"]:
    add_bullet(doc, item)
add_caption(doc, "图 1  V&V-first PIO 资格化流程")
add_picture_with_alt(doc, flow_path, "V&V-first PIO qualification pipeline from reference construction to route closure")

doc.add_heading("2. PIO 理论合同", level=1)
doc.add_heading("2.1 增量与回退", level=2)
add_body(doc, "PIO 输出的是对冻结 SPH RHS 的增量，而不是替代 baseline。zero correction 时必须逐状态 bitwise 恢复 baseline。")
add_equation(doc, "a_hybrid = a_SPH + Δâ,    Δâ = 0 ⇒ a_hybrid = a_SPH", "2")
doc.add_heading("2.2 Pair-force 守恒形式", level=2)
add_body(doc, "优先 scope 是每个无向 pair 恰计一次的互易 pair-force correction。对封闭周期系统，pair exchange antisymmetry 是离散线性动量守恒的充分构造条件。")
add_equation(doc, "ΔF_ji = −ΔF_ij,    Δâ_i = m_i⁻¹ Σ_j ΔF_ij", "3")
add_equation(doc, "Σ_i m_i Δâ_i = Σ_{\u007bi,j\u007d}(ΔF_ij + ΔF_ji) = 0", "4")
doc.add_heading("2.3 对称性与 feature contract", level=2)
for item in ["particle permutation equivariance；", "translation invariance 与 Galilean invariance；", "periodic minimum-image consistency；", "rotation/reflection O(2) equivariance；", "pair exchange symmetry；", "禁止 target、reference、role、ID、regularity metric 作为输入。"]:
    add_bullet(doc, item)
add_callout(doc, "解释边界", "理论充分构造与 metamorphic PASS 只证明 architecture validity，不证明可训练性、泛化或动态稳定性。")

page_break(doc)
doc.add_heading("3. Target 与 reference qualification", level=1)
doc.add_heading("3.1 Reference hierarchy", level=2)
add_table(doc, ["类别", "作用", "资格边界"], [["R1 continuum-compatible", "MMS/解析验证", "需 model-form alignment；不自动是训练 target"], ["R2 semidiscrete-qualified", "时间误差与状态漂移隔离", "默认 diagnostic"], ["R3 independent benchmark", "独立 validation", "禁止训练"], ["RX model-form-misaligned", "反例与诊断", "硬禁止训练"]], [1.7, 2.4, 2.4], font_size=9.1)
doc.add_heading("3.2 从 R2S 诊断到独立 reference", level=2)
add_body(doc, "Stage 02F/G 表明 same-state R2S 可产生非零空间候选，但 resolution smoothness 与 reconstruction bias 仍不足以资格化。Stage 02H 因而引入 Fourier 与 analytic 两个独立 reference；二者在冻结 periodic-vortex case matrix 上通过 normalized L2/Linf、pattern cosine、uncertainty 与 determinism gates。")
doc.add_heading("3.3 Spatial target pool", level=2)
add_body(doc, "Stage 02I materialized 七个 candidate_discretization_target。五个 regular cases 满足 pair-force global residual；两个 jitter cases 保留为 node-residual-only。Stage 02I-R 选择 PAIR_ONLY_REGULAR_SCOPE，没有通过均值扣除或 projection 修改 target。")
add_equation(doc, "Q_L2 = RMS(Δâ − Δa) / (RMS(Δa) + ε_metric)", "5")
add_table(doc, ["阶段", "Reference/target 结果", "硬边界"], [["02E", "8 nonzero candidates; 0 qualified", "temporal/reference derivative dominated"], ["02F/G", "5 R2S candidates; diagnostic closure", "bias/resolution evidence insufficient"], ["02H", "Fourier–analytic independent agreement PASS", "scope-limited"], ["02I/I-R", "7 attributed; 5 pair-only", "2 jitter excluded from pair labels"]], [1.2, 2.8, 2.5], font_size=9.0)

doc.add_heading("4. Conservation-compatible scope", level=1)
add_body(doc, "守恒 scope 的核心是区分 target scientific attribution 与 pair-force representability。Stage 02I 的 jitter target 即使 reference/attribution PASS，也因 nonzero normalized global-force residual 不能进入 pair-force label scope。Stage 02I-R 进一步分解 pressure、viscosity 和 total force，并确认 regular subset 的 roundoff-scale closure。")
for item in ["pair residual ≤ 1e−10；", "normalized total force ≤ 1e−10；", "topology必须 reciprocal、无 duplicate/unexpected edges；", "不得以 projection 写回 target；", "K1/K2 的 torque/power 在非中心力下保持 diagnostic。"]:
    add_bullet(doc, item)
add_callout(doc, "失败语义", "conservation-scope failure 只排除不兼容 target；它不等同于 reference construction failure，也不等同于 architecture failure。", RED)

page_break(doc)
doc.add_heading("5. Dataset construction and leakage control", level=1)
doc.add_heading("5.1 早期 dataset 路线", level=2)
add_body(doc, "Stage 02J 的五个 regular records 虽通过 schema/canonical/QC，却属于单一 leakage component，无法形成合法 split。Stage 02J-R 新增三条预注册 lineage，但十五个 candidate 因 regularity attribution 仅 5/6 diagnostic 而未物化。两次 NOT READY 都被永久保留。")
doc.add_heading("5.2 最终 blind multifamily dataset", level=2)
add_body(doc, "Stage 02J-W 在 regularity diagnostic-only 的独立 eligibility contract 下，按冻结 generator source/config、formula 与 seed 单次物化四个 blind families。20/20 records 通过 reference、target core、pair-only conservation、canonical QC 与 eligibility；leakage graph 恰有四个 family components。")
add_table(doc, ["属性", "冻结结果", "边界"], [["记录数", "20 full graphs", "不是大规模 dataset"], ["Family components", "4 lineage-disconnected", "family是统计/泄漏单元"], ["Split", "10 train / 5 validation / 5 test", "无跨 split lineage"], ["Normalization", "10 train graphs only", "validation/test未参与拟合"], ["Regularity", "diagnostic_only", "不再是 eligibility hard gate"]], [1.6, 2.2, 2.7], font_size=9.2)

doc.add_heading("6. Regularity-hard-gate route and termination", level=1)
add_body(doc, "Regularity-hard-gate 路线经历 v0.1、v0.2、单一 v0.3 candidate 与最后的 v0.4 candidate。每个版本都冻结候选与阈值，并保留失败机制；任何后续状态都未删除历史 false positive 或 invariance failure。")
add_table(doc, ["版本/阶段", "主要证据", "失败机制", "状态"], [["v0.1 / 02J-R", "15 candidates, reference/conservation PASS", "permuted-null ratio gate", "NOT READY"], ["v0.2 / 02J-S", "structured paths + invariance PASS", "negative-control false positives", "NOT READY"], ["v0.3 / 02J-T", "magnitude-direction conjunction", "CROSSMODE N12 magnitude", "NOT QUALIFIED"], ["v0.4 / 02J-V", "necessity/Bonferroni", "9/192 invariance rows", "ROUTE TERMINATED"]], [1.3, 2.3, 2.2, 1.2], font_size=8.7)
add_body(doc, "Stage 02J-W 的 dataset READY 是另一条 eligibility contract 的结果：regularity只作为 diagnostic registry。它不覆盖 J/J-R/J-S/J-T/J-V 的失败，也不重新打开 regularity-hard-gate route。")

doc.add_heading("7. Blind multifamily dataset", level=1)
add_body(doc, "最终 collection identity 为 blind_multifamily_pair_scope_v1_0。四个 family 在 initial-condition、solution、source 与 derivative lineage 上独立；共同的 EOS、SPH/Fourier implementation、schema 和 serializer 被正确视为 infrastructure，而不是数据 ancestry。")
for item in ["所有公式、root seed、split role 在 target evaluation 前冻结；", "reference 20/20 PASS；target core 20/20 PASS；", "resolution/support consistency 按 family 评价；", "pair-only conservation 20/20 PASS，无 target modification；", "train-only graph-balanced normalization；", "jitter 与 R3 evidence 保持 diagnostic/independent-validation isolation。"]:
    add_number(doc, item)
add_caption(doc, "图 2  Stage 02 完整状态时间线")
add_picture_with_alt(doc, timeline_path, "Timeline of 22 Stage 02 states with preserved failures and terminal route state")

page_break(doc)
doc.add_heading("8. Pair-force architecture", level=1)
add_body(doc, "Stage 02K 比较 K0 central diagnostic、K1 conservative pair MLP、K2 reciprocal pair attention PIO 与 KNEG directed-softmax negative control。Architecture hash 在 target array access 前冻结。")
add_table(doc, ["Arm", "形式", "结构资格", "解释"], [["K0", "central pair MLP", "diagnostic", "central representability/torque baseline"], ["K1", "non-attention pair MLP", "QUALIFIED", "antisymmetry/momentum/O(2)/periodicity PASS"], ["K2", "reciprocal pair attention", "QUALIFIED", "same hard gates PASS"], ["KNEG", "directed softmax", "negative control PASS", "暴露非互易 pair/conservation failure"]], [0.9, 1.8, 1.5, 2.3], font_size=8.8)
add_body(doc, "K1/K2 的 qualified 只授权 Stage 02L 做训练协议预注册，不授权正式训练；Stage 02L 随后仅通过独立阶段授权 Stage 02M。K2 的 attention 结构没有建立相对 K1 的必要性或优越性。")

doc.add_heading("9. Static fitting protocol v0.1", level=1)
add_body(doc, "Stage 02L 冻结 K0/K1/K2 × seeds 的九运行矩阵、CPU float64、full-graph balanced loss、AdamW、warmup/cosine、20-update validation/checkpoint、early stopping、一次性 test release 和 A–E success gates。")
add_equation(doc, "L = (1/10) Σ_g mean_i ‖(Δâ_i − Δa_i)/a₀‖²", "6")
add_body(doc, "Stage 02M 完成九个 terminal runs、validation-only checkpoint selection、一次性 sealed test、postfit conservation/symmetry 和 resource audit。尽管结构与部分 transfer evidence 完整，K1/K2 未满足冻结 A–E；最终状态 STATIC_PAIR_FORCE_FITTING_NOT_QUALIFIED。")
add_callout(doc, "负结果保留", "test 或 validation 的局部 PASS 不覆盖 train-fit failure；历史 test 已消费，后续只能作为描述性历史证据。", RED)

doc.add_heading("10. Failure attribution", level=1)
add_body(doc, "Stage 02M-R 未训练、未读取新 test，而是对 frozen v0.1 histories/checkpoints 做 post-hoc 诊断。证据显示 a₀=400 的 loss scaling、Adam epsilon=1e−8 与 weight_decay=1e−6 共同产生不利的优化 conditioning；因此唯一状态为 STATIC_FITTING_FAILURE_ATTRIBUTED_OPTIMIZATION_CONDITIONING。")
add_body(doc, "该归因是 contribution-level diagnostic，不等价于“唯一科学根因已证明”，也不保证修改 conditioning 后可达到 train-fit。feature identifiability audit 未发现硬矛盾，但未证明 mapping 可识别。")

doc.add_heading("11. Static fitting protocol v0.2", level=1)
doc.add_heading("11.1 Prospective changes", level=2)
add_body(doc, "Stage 02M-P 只做前瞻设计：用10个 train graphs 冻结 a_sup=0.392220124168075 m s−2，AdamW epsilon 改为1e−12、weight decay改为0；同时生成新的 validation/test blind families 与 v1.1 collection。architecture、输入 normalization、budget、success gates 与 test seal 均保持可追踪身份。")
add_equation(doc, "L_g = mean_i ‖(Δâ_i/a_sup) − (Δa_i/a_sup)‖²", "7")
doc.add_heading("11.2 Formal execution and conditioning", level=2)
add_body(doc, "Stage 02M-Q 执行唯一九运行重试。每个 run 保存 updates 0、1、10、50、100、selected、terminal 的 conditioning snapshots；九份 histories 全部 PASS。K0/K1 各三个 seeds 达到1000 updates；K2 为700、740、1000，其中前两项按预注册 patience early stop。")
add_table(doc, ["Architecture", "Train gate B", "Validation gate C", "Test gate D", "Conservation E", "A–E"], [["K0 diagnostic", "0/3", "3/3", "3/3", "PASS", "FAIL"], ["K1", "0/3", "3/3", "3/3", "PASS", "FAIL"], ["K2", "1/3", "3/3", "3/3", "PASS", "FAIL"]], [1.3, 1.1, 1.2, 1.0, 1.1, 0.8], font_size=8.8)
add_body(doc, "v0.2 的 transfer gates 与 conditioning 改善只能做 descriptive protocol comparison。v0.1/v0.2 validation/test families 不同，不能当 paired benchmark；最优 K2 seed 不能替代 2-of-3 seed rule。")

doc.add_heading("12. Final route termination", level=1)
add_callout(doc, "Terminal verdict", "STATIC_PAIR_FORCE_FITTING_V02_NOT_QUALIFIED", RED)
add_body(doc, "终止条件已经满足：两次正式静态协议完整执行且均未资格化；第二次是 optimization-conditioning 路线的唯一允许重试。故 static_pio_learning_route_terminated=true，training_protocol_v03_permitted=false，Stage02N_authorized=false。")
for item in ["不得增加 seed；", "不得修改 loss scale 或 optimizer；", "不得继续生成 blind families 改善结果；", "不得进入 solver integration；", "未来任何新假设必须作为全新 Stage 03。"]:
    add_bullet(doc, item)

doc.add_heading("13. Scientific conclusions", level=1)
for claim in claims["supported_claims"]:
    add_body(doc, "支持：" + claim["allowed_wording"])
add_body(doc, "最重要的科学结论不是 attention 成功，而是 architecture correctness、reference/dataset qualification 与 static learnability 是不同层级的命题。Stage 02 同时获得了强结构证据和完整负学习证据，因此能支持可证伪 qualification framework 的方法学主张。")
add_caption(doc, "图 3  Supported / conditional / unsupported claims 边界")
add_picture_with_alt(doc, boundary_path, "Supported, conditional, and unsupported scientific claim boundary")

doc.add_heading("14. Unsupported claims", level=1)
add_table(doc, ["Unsupported claim", "允许写法", "禁止写法"], [[row["claim"], row["allowed_wording"], row["prohibited_wording"]] for row in claims["unsupported_claims"]], [1.8, 2.5, 2.2], font_size=8.3)
add_callout(doc, "动态语义", "rollout 和 solver-in-the-loop 是 NOT AUTHORIZED / NOT EXECUTED；禁止写成 rollout failed、solver unstable 或 solver accelerated。", RED)

doc.add_heading("15. Future research branches", level=1)
add_body(doc, "以下分支只构成 decision design，不在 Stage 02M-S 执行。每个分支必须以全新 Stage 03 建立新的 hypothesis、reference/V&V、freeze、dataset 和 success contract。")
for row in future["branches"]:
    doc.add_heading(f"15.{row['branch']} Branch {row['branch']}：{row['title']}", level=2)
    add_body(doc, "科学假设：" + row["hypothesis"])
    add_body(doc, "阶段要求：" + row["required_stage"] + "；允许设计：" + "、".join(row["allowed_design"]) + "。本阶段执行：否。")
doc.add_heading("15.5 推荐优先级", level=2)
add_body(doc, "优先 Branch 1，将 Stage 02 组织为 V&V/qualification framework + falsified static learning 的论文；其次 Branch 2 与 Branch 4，用低维 identifiable target 和非神经 conservative baseline 建立机制对照；Branch 3 的 trajectory dataset 代价与 reference 重建需求最高。")

doc.add_heading("16. Artifact and hash index", level=1)
add_body(doc, f"Stage 02M-S freeze manifest 共记录 {freeze['historical_file_count']} 个历史文件。下面列出主索引；完整逐文件 SHA-256 见 08_route_closure/freeze/stage02ms_historical_freeze_manifest.json。")
index_rows = []
for row in ledger["rows"]:
    index_rows.append([row["stage"], row["unique_status"], row["principal_evidence"]["artifact"], row["historical_hash"][:24] + "…"])
add_table(doc, ["Stage", "Status", "Principal artifact", "SHA-256 prefix"], index_rows[:11], [0.9, 2.0, 2.7, 0.9], font_size=7.7)
page_break(doc)
doc.add_heading("16.1 Artifact and hash index（续）", level=2)
add_table(doc, ["Stage", "Status", "Principal artifact", "SHA-256 prefix"], index_rows[11:], [0.9, 2.0, 2.7, 0.9], font_size=7.7)

page_break(doc)
doc.add_heading("附录 A：完整 Stage 02 状态账本", level=1)
add_body(doc, "本表按时间顺序保留22个状态；superseded 对所有行均为 false。")
ledger_rows = [[row["order"], row["stage"], row["unique_status"], row["execution_count"], row["training_runs"], row["optimizer_steps"], row["principal_blocker"]] for row in ledger["rows"]]
add_table(doc, ["#", "Stage", "Unique status", "Exec", "Runs", "Steps", "Blocker"], ledger_rows[:11], [0.35, 0.75, 1.9, 0.55, 0.45, 0.65, 1.85], font_size=7.6)
page_break(doc)
doc.add_heading("附录 A.1：完整 Stage 02 状态账本（续）", level=2)
add_table(doc, ["#", "Stage", "Unique status", "Exec", "Runs", "Steps", "Blocker"], ledger_rows[11:], [0.35, 0.75, 1.9, 0.55, 0.45, 0.65, 1.85], font_size=7.6)

page_break(doc)
doc.add_heading("附录 B：证据矩阵", level=1)
add_table(doc, ["ID", "Class", "Level", "Claim", "Limitation", "Negative"], [[row["id"], row["claim_class"], row["evidence_level"], row["claim"], row["limitations"], "yes" if row["negative_result"] else "no"] for row in matrix["rows"]], [0.4, 1.0, 1.1, 1.7, 1.9, 0.4], font_size=7.6)

page_break(doc)
doc.add_heading("附录 C：失败分类与状态语义", level=1)
add_table(doc, ["#", "Failure class", "Definition", "Recorded instances"], [[row["id"], row["name"], row["definition"], "; ".join(f"{x['stage']}:{x['state']}" for x in row["instances"])] for row in taxonomy["classes"]], [0.4, 1.55, 2.4, 2.15], font_size=8.0)

page_break(doc)
doc.add_heading("附录 D：论文可行性与图表包", level=1)
add_table(doc, ["Paper", "Direction", "Readiness", "CMAME now", "Fatal weakness"], [[row["paper"], row["direction"], row["manuscript_readiness"], row["current_CMAME_target_defensible"], row["fatal_weakness"]] for row in assessment["papers"]], [0.65, 2.0, 1.25, 0.7, 1.9], font_size=8.1)
add_body(doc, "工作标题：" + assessment["working_title"])
add_body(doc, "CMAME 官方 scope 来源：" + assessment["CMAME_scope_source"])
add_table(doc, ["Fig.", "Title", "Form", "Integrity rule"], [[row["figure"], row["title"], row["form"], row["integrity_rule"]] for row in figure_plan["figures"]], [0.45, 2.2, 0.9, 2.95], font_size=7.9)
add_body(doc, "计划表格：" + "；".join(row["title"] for row in figure_plan["tables"]) + "。")

page_break(doc)
doc.add_heading("附录 E：Stage 02M-Q selected checkpoint index", level=1)
checkpoint_rows = [[row["run_id"], row["selected_checkpoint"].split("/")[-1], row["expected_sha256"][:28] + "…", row["status"]] for row in freeze["selected_checkpoints"]]
add_table(doc, ["Run", "Checkpoint", "SHA-256 prefix", "Identity"], checkpoint_rows, [1.45, 1.15, 2.75, 1.15], font_size=8.2)
add_body(doc, "所有 selected checkpoints 在 test release 前被冻结；每个 checkpoint 在新 sealed test 上评价一次，test 后 optimizer steps=0，checkpoint changes=0。")

doc.add_heading("记录闭包声明", level=1)
add_callout(doc, "Stage 02M-S", "STAGE02_ROUTE_CLOSED_PUBLICATION_BOUNDARY_COMPLETE 的判定须以最终 render QA、历史复核和 closure manifest 为准。本研究记录本身不授权新训练、新测试、rollout 或 solver-in-the-loop。", NAVY)

doc.core_properties.title = "Stage 02 Research Record"
doc.core_properties.subject = "SPH-PIO-PoC route closure and publication boundary"
doc.core_properties.author = "SPH-PIO-PoC Research Record"
doc.core_properties.keywords = "SPH, PIO, verification, qualification, conservative learning, negative result"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(json.dumps({"output": str(OUT), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "figures": 3}, ensure_ascii=False))
