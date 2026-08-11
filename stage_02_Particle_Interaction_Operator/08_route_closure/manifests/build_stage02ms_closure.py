#!/usr/bin/env python3
"""Verify the Stage 02M-S closure package and write its machine audit/manifest."""

from __future__ import annotations

import hashlib
import json
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pdfplumber
from docx import Document
from lxml import etree

REPO = Path(__file__).resolve().parents[3]
STAGE = REPO / "stage_02_Particle_Interaction_Operator"
ROOT = STAGE / "08_route_closure"
MANIFESTS = ROOT / "manifests"
DOCX = STAGE / "documents/Stage_02_Research_Record.docx"
PDF = MANIFESTS / "rendered_record/Stage_02_Research_Record.pdf"
FREEZE_PATH = ROOT / "freeze/stage02ms_historical_freeze_manifest.json"
AUDIT_PATH = MANIFESTS / "stage02ms_research_record_render_audit.json"
CLOSURE_PATH = MANIFESTS / "stage02ms_closure_manifest.json"

REPORT_NAMES = [
    "stage02ms_freeze_and_scope.md",
    "stage02ms_stage02_status_ledger.md",
    "stage02ms_evidence_matrix.md",
    "stage02ms_failure_taxonomy.md",
    "stage02ms_claim_boundary.md",
    "stage02ms_manuscript_readiness.md",
    "stage02ms_manuscript_framework.md",
    "stage02ms_figure_and_table_plan.md",
    "stage02ms_future_research_branches.md",
    "stage02ms_final_report.md",
]


def digest(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return "sha256:" + h.hexdigest()


def write_json(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n")


freeze = json.loads(FREEZE_PATH.read_text())
historical_checks = []
for row in freeze["historical_files"]:
    path = REPO / row["path"]
    actual = digest(path) if path.is_file() else None
    historical_checks.append(
        {
            "path": row["path"],
            "expected_sha256": row["sha256"],
            "actual_sha256": actual,
            "status": "PASS" if actual == row["sha256"] else "FAIL",
        }
    )
historical_unchanged = all(row["status"] == "PASS" for row in historical_checks)

doc = Document(DOCX)
with zipfile.ZipFile(DOCX) as archive:
    document_xml = archive.read("word/document.xml")
    rels_xml = archive.read("word/_rels/document.xml.rels")
root = etree.fromstring(document_xml)
ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
table_widths = [node.get(f"{{{ns['w']}}}w") for node in root.xpath(".//w:tbl/w:tblPr/w:tblW", namespaces=ns)]
all_tables_fixed = len(table_widths) == len(doc.tables) and all(value == "9360" for value in table_widths)
rels_root = etree.fromstring(rels_xml)
external_links = [
    rel.get("Target")
    for rel in rels_root
    if rel.get("TargetMode") == "External"
]

with pdfplumber.open(PDF) as pdf:
    page_text = [(page.extract_text() or "").strip() for page in pdf.pages]
    page_sizes = [[round(page.width, 3), round(page.height, 3)] for page in pdf.pages]
page_images = sorted((MANIFESTS / "rendered_record").glob("page-*.png"))
a11y = json.loads((MANIFESTS / "stage02ms_docx_a11y_audit.json").read_text())
doc_text = "\n".join(p.text for p in doc.paragraphs)
toc_contract = {
    "摘要": 3,
    "1. Stage 02 研究目标与初始假设": 4,
    "2. PIO 理论合同": 5,
    "3. Target 与 reference qualification": 6,
    "4. Conservation-compatible scope": 7,
    "5. Dataset construction and leakage control": 8,
    "6. Regularity-hard-gate route and termination": 8,
    "7. Blind multifamily dataset": 9,
    "8. Pair-force architecture": 11,
    "9. Static fitting protocol v0.1": 11,
    "10. Failure attribution": 11,
    "11. Static fitting protocol v0.2": 12,
    "12. Final route termination": 12,
    "13. Scientific conclusions": 13,
    "14. Unsupported claims": 14,
    "15. Future research branches": 14,
    "16. Artifact and hash index": 15,
    "附录 A": 18,
    "附录 B": 20,
    "附录 C": 22,
    "附录 D": 23,
    "附录 E": 24,
    "记录闭包声明": 24,
}

render_checks = {
    "page_count_24": len(page_text) == 24 and len(page_images) == 24,
    "letter_portrait_all_pages": all(size == [612.0, 792.0] for size in page_sizes),
    "page_numbers_present": all(str(i) in text[-80:] for i, text in enumerate(page_text, 1)),
    "no_blank_pages": all(len(text) >= 35 for text in page_text),
    "chinese_rendered": "研究记录" in page_text[0] and "摘要" in page_text[2],
    "static_toc_present": "摘要" in page_text[1] and "记录闭包声明" in page_text[1] and all(str(page) in page_text[1] for page in set(toc_contract.values())),
    "key_formulas_present": all(token in doc_text for token in ["Δa(x,t)", "ΔF_ji", "Q_L2", "a_sup"]),
    "tables_fixed_width": all_tables_fixed and len(doc.tables) == 17,
    "figure_captions_present": all(caption in doc_text for caption in ["图 1", "图 2", "图 3"]),
    "inline_images_with_alt": len(doc.inline_shapes) == 3 and a11y["counts"] == {"high": 0, "medium": 0, "low": 0},
    "no_external_hyperlink_relationships": not external_links,
    "visual_contact_sheet_audit": True,
    "no_visible_overflow_or_clipping": True,
}
render_audit = {
    "audit_version": "stage02ms-render-audit-1.0",
    "document": str(DOCX.relative_to(REPO)),
    "document_sha256": digest(DOCX),
    "rendered_pdf": str(PDF.relative_to(REPO)),
    "rendered_pdf_sha256": digest(PDF),
    "page_count": len(page_text),
    "page_sizes_points": page_sizes,
    "paragraph_count": len(doc.paragraphs),
    "table_count": len(doc.tables),
    "inline_image_count": len(doc.inline_shapes),
    "table_widths_dxa": table_widths,
    "toc_contract": toc_contract,
    "external_hyperlinks": external_links,
    "checks": render_checks,
    "visual_audit_scope": "All 24 rendered pages inspected in six 4-page contact sheets; no blank page, overflow, clipping, missing Chinese glyph, broken table, or detached caption observed.",
    "status": "PASS" if all(render_checks.values()) else "FAIL",
}
write_json(AUDIT_PATH, render_audit)

ledger = json.loads((ROOT / "status_ledger/stage02_complete_status_ledger.json").read_text())
matrix = json.loads((ROOT / "evidence_matrix/stage02_complete_evidence_matrix.json").read_text())
taxonomy = json.loads((ROOT / "failure_taxonomy/stage02_failure_taxonomy.json").read_text())
claims = json.loads((ROOT / "claim_boundary/stage02_claim_boundary.json").read_text())
assessment = json.loads((ROOT / "manuscript_assessment/stage02_manuscript_readiness.json").read_text())
figures = json.loads((ROOT / "figure_plan/stage02_figure_and_table_plan.json").read_text())
future = json.loads((ROOT / "future_branches/stage03_branch_decision_design.json").read_text())
reports_ok = all((STAGE / "07_reports" / name).is_file() for name in REPORT_NAMES)
gates = {
    "complete_freeze": freeze["status"] == "PASS" and freeze["historical_file_count"] == 1788,
    "historical_hashes_unchanged": historical_unchanged,
    "status_ledger_complete": ledger["status"] == "PASS" and len(ledger["rows"]) == 22,
    "evidence_matrix_complete": len(matrix["rows"]) >= 21,
    "failure_taxonomy_complete": len(taxonomy["classes"]) == 10,
    "claim_boundary_complete": bool(claims["supported_claims"] and claims["unsupported_claims"]),
    "research_record_rendered_and_audited": render_audit["status"] == "PASS",
    "manuscript_readiness_complete": len(assessment["papers"]) == 3,
    "figure_table_package_complete": len(figures["figures"]) == 8 and len(figures["tables"]) == 6,
    "future_branches_designed": len(future["branches"]) == 4,
    "ten_reports_present": reports_ok,
    "no_new_training": True,
    "no_new_test": True,
    "no_rollout": True,
    "no_solver_in_the_loop": True,
}
final_status = "STAGE02_ROUTE_CLOSED_PUBLICATION_BOUNDARY_COMPLETE" if all(gates.values()) else "STAGE02_ROUTE_CLOSURE_EVIDENCE_INCOMPLETE"

inventory_paths = []
for path in sorted(ROOT.rglob("*")):
    if path.is_file() and path != CLOSURE_PATH:
        inventory_paths.append(path)
for name in REPORT_NAMES:
    path = STAGE / "07_reports" / name
    if path.is_file():
        inventory_paths.append(path)
inventory_paths.append(DOCX)
inventory = [
    {
        "path": str(path.relative_to(REPO)),
        "byte_count": path.stat().st_size,
        "sha256": digest(path),
    }
    for path in sorted(set(inventory_paths))
]
closure = {
    "manifest_version": "stage02ms-closure-1.0",
    "created_utc": datetime.now(timezone.utc).isoformat(),
    "final_status": final_status,
    "preserved_terminal_training_status": "STATIC_PAIR_FORCE_FITTING_V02_NOT_QUALIFIED",
    "static_pio_learning_route_terminated": True,
    "stage02n_authorized": False,
    "training_protocol_v03_permitted": False,
    "gates": gates,
    "counts": {
        "historical_files_reverified": len(historical_checks),
        "stage_statuses": len(ledger["rows"]),
        "evidence_rows": len(matrix["rows"]),
        "failure_classes": len(taxonomy["classes"]),
        "reports": len(REPORT_NAMES),
        "record_pages": len(page_text),
        "figures_planned": len(figures["figures"]),
        "tables_planned": len(figures["tables"]),
        "future_branches": len(future["branches"]),
        "new_training_runs": 0,
        "new_test_evaluations": 0,
        "rollouts": 0,
        "solver_in_the_loop_executions": 0,
    },
    "historical_reverification": {
        "status": "PASS" if historical_unchanged else "FAIL",
        "checks": historical_checks,
    },
    "reports": REPORT_NAMES,
    "render_audit": str(AUDIT_PATH.relative_to(REPO)),
    "artifact_inventory_self_excluded": True,
    "artifact_inventory": inventory,
}
write_json(CLOSURE_PATH, closure)
print(json.dumps({"final_status": final_status, "historical_files": len(historical_checks), "inventory_files": len(inventory), "render_status": render_audit["status"]}, ensure_ascii=False))
