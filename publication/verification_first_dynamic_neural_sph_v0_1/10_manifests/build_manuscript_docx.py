#!/usr/bin/env python3
"""Build the Chinese Publication P1 manuscript DOCX from the evidence-locked Markdown."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PUB = Path(__file__).resolve().parents[1]
SOURCE = PUB / "03_manuscript_cn/manuscript_cn_v0_1.md"
OUTPUT = PUB / "03_manuscript_cn/manuscript_cn_v0_1.docx"
TOC_MAP = PUB / "10_manifests/toc_pages.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "667085"
LIGHT = "F4F6F9"
TABLE_FILL = "E8EEF5"
RED_FILL = "FDECEC"
GOLD_FILL = "FFF5D6"
GREEN_FILL = "EAF5EF"
FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Heiti SC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia=FONT_EAST_ASIA):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style(style, size, color, bold, before, after, line, alignment=None):
    style.font.name = FONT_LATIN
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if alignment is not None:
        pf.alignment = alignment


def add_field(paragraph, instruction: str, fallback: str = "1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = fallback
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str, color=BLUE):
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN); rfonts.set(qn("w:hAnsi"), FONT_LATIN); rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    rpr.append(rfonts)
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "none"); rpr.append(u)
    run.append(rpr)
    node = OxmlElement("w:t"); node.text = text; run.append(node)
    hyperlink.append(run); paragraph._p.append(hyperlink)


def paragraph_border(paragraph, side="left", color=BLUE, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr"); ppr.append(pbdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single"); border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space); border.set(qn("w:color"), color)
    pbdr.append(border)


def paragraph_shading(paragraph, fill=LIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); ppr.append(shd)


def add_inline_markup(paragraph, text: str):
    # Keep REF-TODO and evidence statuses visible; remove only source-side claim comments.
    text = re.sub(r"\s*<!--\s*CLAIM:[^>]+-->\s*", "", text)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); set_run_font(run, size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part); set_run_font(run)


def fixed_table(doc: Document, rows: list[list[str]]):
    cols = len(rows[0])
    patterns = {
        3: [1900, 1100, 6360],
        4: [1350, 2500, 1500, 4010],
    }
    widths = patterns.get(cols, [9360 // cols] * cols)
    widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360"); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for r_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        if r_idx == 0:
            flag = OxmlElement("w:tblHeader"); flag.set(qn("w:val"), "true"); tr_pr.append(flag)
        for c_idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[c_idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[c_idx])); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0: set_cell_shading(cell, TABLE_FILL)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markup(p, rows[r_idx][c_idx])
            for run in p.runs:
                set_run_font(run, size=8.5, bold=(r_idx == 0), color=(NAVY if r_idx == 0 else "1F2937"))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)
    return table


def parse_markdown_table(lines: list[str], index: int):
    rows = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def heading_anchor(text: str) -> str:
    if text == "摘要": return "abstract"
    if text == "关键词": return "keywords"
    if re.match(r"^\d+\.", text): return "sec" + text.split(".")[0]
    return {
        "Data availability": "data_availability",
        "Code availability": "code_availability",
        "Author contributions": "author_contributions",
        "Conflict of interest": "conflict_of_interest",
        "References": "references",
    }.get(text, "h_" + str(abs(hash(text))))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

# narrative_proposal token map.
styles = doc.styles
set_style(styles["Normal"], 11, "1F2937", False, 0, 8, 1.333, WD_ALIGN_PARAGRAPH.JUSTIFY)
set_style(styles["Heading 1"], 16, BLUE, True, 18, 10, 1.05, WD_ALIGN_PARAGRAPH.LEFT)
set_style(styles["Heading 2"], 13, BLUE, True, 12, 6, 1.05, WD_ALIGN_PARAGRAPH.LEFT)
set_style(styles["Heading 3"], 12, DARK_BLUE, True, 8, 4, 1.05, WD_ALIGN_PARAGRAPH.LEFT)
styles["Heading 1"].paragraph_format.keep_with_next = True
styles["Heading 2"].paragraph_format.keep_with_next = True

# Running header/footer.
header = section.header.paragraphs[0]
header.text = "SPH-PIO-PoC  |  Evidence-Locked Manuscript v0.1"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in header.runs: set_run_font(run, size=8.5, color=MUTED)
paragraph_border(header, side="bottom", color="D7DEE8", size="6", space="5")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("Publication Track P1  |  "); set_run_font(r, size=9, color=MUTED)
add_field(footer, "PAGE", "1")

# Editorial cover.
for _ in range(5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PUBLICATION TRACK P1"); set_run_font(r, size=10, bold=True, color="B47D18")
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("守恒型动态神经–SPH求解器的\n验证优先开发"); set_run_font(r, size=23, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("零修正等价、拓扑事件与多步梯度资格边界"); set_run_font(r, size=17, color=DARK_BLUE)
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Verification-first development of a conservative dynamic neural–SPH solver"); set_run_font(r, size=12.5, italic=True, color=MUTED)
p.paragraph_format.space_after = Pt(36)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("中文稿 v0.1  |  Evidence-locked  |  非投稿定稿"); set_run_font(r, size=11, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Stage 03D: NOT_QUALIFIED  |  Training / rollout: NOT EXECUTED  |  Stage 03E: false"); set_run_font(r, size=9.5, color="9B1C1C")
paragraph_shading(p, RED_FILL); paragraph_border(p, side="left", color="9B1C1C", size="18", space="8")
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
doc.add_page_break()

# Static TOC (top-level only), with internal links.
toc_pages = json.loads(TOC_MAP.read_text()) if TOC_MAP.is_file() else {}
toc_entries = [
    ("摘要", "abstract"), ("关键词", "keywords"),
    *[(f"{i}. " + ["引言","控制方程与模型形式","验证与资格框架","动态参考轨迹","动态求解器实现","结构验证","多步可微性资格","拓扑事件资格","讨论","结论"][i-1], f"sec{i}") for i in range(1,11)],
    ("Data availability", "data_availability"), ("Code availability", "code_availability"),
    ("Author contributions", "author_contributions"), ("Conflict of interest", "conflict_of_interest"), ("References", "references"),
]
p = doc.add_paragraph("目录", style="Heading 1"); add_bookmark(p, "toc", 1)
for label, anchor in toc_entries:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    add_internal_link(p, label, anchor)
    dots = p.add_run(" ·" * max(4, 35 - len(label))); set_run_font(dots, size=9.5, color="98A2B3")
    page = p.add_run(" " + str(toc_pages.get(anchor, "—"))); set_run_font(page, size=10, bold=True, color=NAVY)
p = doc.add_paragraph("目录为headless-safe静态索引；条目包含内部链接，页码由最终渲染校正。")
for run in p.runs: set_run_font(run, size=8.5, italic=True, color=MUTED)
doc.add_page_break()

lines = SOURCE.read_text().splitlines()
# Skip cover source lines until the source Abstract heading.
start = next(i for i, line in enumerate(lines) if line.strip() == "## 摘要")
lines = lines[start:]
bookmark_id = 10
i = 0
buffer: list[str] = []


def flush_buffer():
    nonlocal_placeholder = None
    if not buffer:
        return
    text = " ".join(item.strip() for item in buffer).strip()
    buffer.clear()
    if not text:
        return
    p = doc.add_paragraph()
    add_inline_markup(p, text)
    p.paragraph_format.first_line_indent = Inches(0.28)
    if any(token in text for token in ("DYNAMIC_", "STAGE02_ROUTE_", "TOPOLOGY_EVENT_COMPONENT_", "V2_QUALIFICATION_", "UNSUPPORTED_DRAFT_STATEMENT")):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


while i < len(lines):
    line = lines[i].rstrip()
    stripped = line.strip()
    if not stripped:
        flush_buffer(); i += 1; continue
    if stripped.startswith("<!--"):
        i += 1; continue
    if stripped.startswith("|"):
        flush_buffer(); rows, i = parse_markdown_table(lines, i); fixed_table(doc, rows); continue
    if stripped.startswith("$$") and stripped.endswith("$$"):
        flush_buffer()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8); p.paragraph_format.keep_together = True
        r = p.add_run(stripped[2:-2].strip()); set_run_font(r, size=11, italic=True, color=NAVY)
        i += 1; continue
    if re.match(r"^\*\*图\d+设计。\*\*", stripped):
        flush_buffer()
        match = re.match(r"^\*\*图(\d+)设计。\*\*\s*(.*)$", stripped)
        number, body = match.group(1), match.group(2)
        p = doc.add_paragraph(); paragraph_shading(p, LIGHT); paragraph_border(p, side="left", color=BLUE, size="18", space="8")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.12); p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4); p.paragraph_format.keep_together = True; p.paragraph_format.keep_with_next = True
        r = p.add_run(f"FIGURE {number} — P1 DETAILED DESIGN\n"); set_run_font(r, size=9, bold=True, color=BLUE)
        r = p.add_run(body); set_run_font(r, size=9.5, color="344054")
        cap = doc.add_paragraph(style="Caption"); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8); cap.paragraph_format.keep_with_next = False
        r = cap.add_run(f"图{number}  P1证据锁定图件设计（最终科研图待P2/Python或R后端确认）"); set_run_font(r, size=9, italic=True, color=MUTED)
        i += 1; continue
    if stripped.startswith("#"):
        flush_buffer()
        level = len(stripped) - len(stripped.lstrip("#"))
        text = stripped[level:].strip()
        # Source uses ## for abstract/keywords but they are top-level front matter.
        effective_level = 1 if text in ("摘要", "关键词") else min(level, 2)
        p = doc.add_paragraph(text, style=f"Heading {effective_level}")
        anchor = heading_anchor(text)
        if effective_level == 1:
            add_bookmark(p, anchor, bookmark_id); bookmark_id += 1
        i += 1; continue
    if re.match(r"^\*\*表\d+", stripped):
        flush_buffer()
        p = doc.add_paragraph(style="Caption"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4); p.paragraph_format.keep_with_next = True
        add_inline_markup(p, stripped)
        for run in p.runs: set_run_font(run, size=9, bold=True, color=NAVY)
        i += 1; continue
    buffer.append(stripped); i += 1
flush_buffer()

# Metadata and update settings.
doc.core_properties.title = "守恒型动态神经–SPH求解器的验证优先开发：中文稿v0.1"
doc.core_properties.subject = "Publication Track P1 evidence-locked manuscript"
doc.core_properties.author = "SPH-PIO-PoC Publication Workflow"
doc.core_properties.keywords = "SPH, verification-first, conservative neural solver, multistep gradient, topology event"
settings = doc.settings.element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields"); settings.append(update)
update.set(qn("w:val"), "true")

doc.save(OUTPUT)
print(json.dumps({"output": str(OUTPUT), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "toc_loaded": TOC_MAP.is_file()}, ensure_ascii=False))
