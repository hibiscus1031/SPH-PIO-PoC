#!/usr/bin/env python3
"""Build P2 manuscript and verified-reference DOCX files from the retained P1 template."""

from __future__ import annotations

import csv
import html
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PUB = Path(__file__).resolve().parents[2]
P2 = PUB / "11_literature_verification"
REFERENCE = PUB / "03_manuscript_cn/manuscript_cn_v0_1.docx"
SOURCE = P2 / "manuscript_revision/manuscript_cn_v0_2_literature_positioned.md"
OUTPUT = P2 / "manuscript_revision/manuscript_cn_v0_2_literature_positioned.docx"
REF_OUTPUT = P2 / "verified_records/references_verified.docx"
CSV = P2 / "verified_records/verified_bibliography.csv"

# LibreOffice's isolated profile does not discover macOS system CJK fonts.
# Source Han Sans CN is supplied through SAL_FONTPATH during render and covers
# both Chinese and Latin, so all font slots deliberately use this family.
LATIN = "Source Han Sans CN"
CJK = "Source Han Sans CN"
BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "17365D"
MUTED = "667085"
LIGHT = "F4F6F9"
TABLE_FILL = "E8EEF5"
RED_FILL = "FDECEC"


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_font(run, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = LATIN
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), LATIN)
    fonts.set(qn("w:hAnsi"), LATIN)
    fonts.set(qn("w:eastAsia"), CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def repair_styles(doc: Document) -> None:
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = doc.styles[name]
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        fonts.set(qn("w:ascii"), LATIN)
        fonts.set(qn("w:hAnsi"), LATIN)
        fonts.set(qn("w:eastAsia"), CJK)


def border(paragraph, side="left", color=BLUE, size="18", space="8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    node = OxmlElement(f"w:{side}")
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), size)
    node.set(qn("w:space"), space)
    node.set(qn("w:color"), color)
    pbdr.append(node)


def shading(paragraph, fill=LIGHT) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    ppr.append(node)


def field(paragraph, instruction: str, fallback="1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = fallback
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])
    set_font(run, 9, color=MUTED)


def add_inline(paragraph, text: str, size=None) -> None:
    text = html.unescape(re.sub(r"\s*<!--\s*CLAIM:[^>]+-->\s*", "", text))
    for part in re.split(r"(\*\*.*?\*\*|`.*?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); set_font(run, size=size or 9.5, color=DARK)
        else:
            run = paragraph.add_run(part); set_font(run, size=size)


def normalize_display_math(text: str) -> str:
    """Avoid renderer-dependent Unicode super/subscript collisions."""
    replacements = {
        "k₁": "k_1", "k₂": "k_2", "tₙ₊₁": "t_(n+1)",
        "Sⁿ⁺¹ᐟ²": "S^(n+1/2)", "Sⁿ⁺¹": "S^(n+1)", "Sⁿ": "S^n",
        "historyⁿ": "history^n", "ρ₀": "rho_0", "ρ_i": "rho_i",
        "Δt": "dt",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}"); mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    patterns = {3: [1900, 1100, 6360], 4: [1350, 2500, 1500, 4010]}
    widths = patterns.get(cols, [9360 // cols] * cols)
    widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    pr = table._tbl.tblPr
    for tag, value, typ in (("tblW", "9360", "dxa"), ("tblInd", "120", "dxa")):
        node = pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}"); pr.append(node)
        node.set(qn("w:w"), value); node.set(qn("w:type"), typ)
    layout = pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol"); node.set(qn("w:w"), str(width)); grid.append(node)
    for ri, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        keep = OxmlElement("w:cantSplit"); keep.set(qn("w:val"), "true"); trpr.append(keep)
        if ri == 0:
            head = OxmlElement("w:tblHeader"); head.set(qn("w:val"), "true"); trpr.append(head)
        for ci, cell in enumerate(row.cells):
            cell.width = Inches(widths[ci] / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW"); tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[ci])); tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), TABLE_FILL); tcpr.append(shade)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, rows[ri][ci], 8.5)
            for run in p.runs:
                if ri == 0:
                    set_font(run, 8.5, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure_header_footer(doc: Document, label: str, footer_label: str) -> None:
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.text = label; header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_font(run, 8.5, color=MUTED)
    border(header, side="bottom", color="D7DEE8", size="6", space="5")
    footer = sec.footer.paragraphs[0]
    footer.text = ""; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(footer_label + "  |  "); set_font(run, 9, color=MUTED)
    field(footer, "PAGE", "1")


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PUBLICATION TRACK P2"); set_font(run, 10, bold=True, color="B47D18"); p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("守恒型动态神经–SPH耦合的\n验证优先资格"); set_font(run, 23, bold=True, color=NAVY); p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("零修正、拓扑事件与多步梯度边界"); set_font(run, 17, color=DARK); p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Verification-first qualification of conservative dynamic neural–SPH coupling"); set_font(run, 12.5, italic=True, color=MUTED); p.paragraph_format.space_after = Pt(36)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中文稿 v0.2  |  Literature-positioned  |  非投稿定稿"); set_font(run, 11, bold=True, color=NAVY); p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Stage 03D: NOT_QUALIFIED  |  Training / rollout: NOT EXECUTED  |  Stage 03E: false"); set_font(run, 9.5, color="9B1C1C")
    shading(p, RED_FILL); border(p, color="9B1C1C"); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_paragraph("目录", style="Heading 1")
    labels = ["摘要", "关键词"] + [f"{i}. {name}" for i, name in enumerate(["引言", "控制方程与模型形式", "验证与资格框架", "动态参考轨迹", "动态求解器实现", "结构验证", "多步可微性资格", "拓扑事件资格", "讨论", "结论"], 1)] + ["Data availability", "Code availability", "Author contributions", "Conflict of interest", "References"]
    for label in labels:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label); set_font(run, 10, color=BLUE)
        dots = p.add_run(" ·" * max(4, 35 - len(label))); set_font(dots, 9.5, color="98A2B3")
    p = doc.add_paragraph("目录为headless-safe静态索引；打开Word后可按标题导航。")
    for run in p.runs:
        set_font(run, 8.5, italic=True, color=MUTED)
    doc.add_page_break()


def parse_table(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_manuscript_body(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## 摘要")
    lines = lines[start:]
    i = 0
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        text = " ".join(x.strip() for x in buffer).strip(); buffer.clear()
        if not text:
            return
        p = doc.add_paragraph(); add_inline(p, text); p.paragraph_format.first_line_indent = Inches(0.28)
        if any(t in text for t in ("DYNAMIC_", "STAGE02_", "TOPOLOGY_EVENT_", "V2_QUALIFICATION_", "NOT_")):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    while i < len(lines):
        s = lines[i].strip()
        if not s:
            flush(); i += 1; continue
        if s.startswith("<!--"):
            i += 1; continue
        if s.startswith("|"):
            flush(); rows, i = parse_table(lines, i); add_table(doc, rows); continue
        if s.startswith("$$") and s.endswith("$$"):
            flush(); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
            run = p.add_run(normalize_display_math(s[2:-2].strip())); set_font(run, 11, italic=True, color=NAVY); i += 1; continue
        if re.match(r"^\*\*图\d+设计。\*\*", s):
            flush(); m = re.match(r"^\*\*图(\d+)设计。\*\*\s*(.*)$", s)
            p = doc.add_paragraph(); shading(p); border(p); p.paragraph_format.keep_together = True
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"FIGURE {m.group(1)} — P2 LITERATURE-POSITIONED DESIGN\n"); set_font(run, 9, bold=True, color=BLUE)
            run = p.add_run(m.group(2)); set_font(run, 9.5, color="344054")
            cap = doc.add_paragraph(style="Caption"); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(f"图{m.group(1)}  P2文献定位图件设计（科研图待选定Python或R后端后制作）"); set_font(run, 9, italic=True, color=MUTED)
            i += 1; continue
        if s.startswith("#"):
            flush(); level = len(s) - len(s.lstrip("#")); text = s[level:].strip()
            level = 1 if text in ("摘要", "关键词") else min(level, 2)
            doc.add_paragraph(text, style=f"Heading {level}"); i += 1; continue
        if re.match(r"^\*\*表\d+", s):
            flush(); p = doc.add_paragraph(style="Caption"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
            add_inline(p, s, 9); i += 1; continue
        buffer.append(s); i += 1
    flush()


def build_manuscript() -> dict:
    doc = Document(REFERENCE)
    clear_body(doc); repair_styles(doc)
    configure_header_footer(doc, "SPH-PIO-PoC  |  Literature-Positioned Manuscript v0.2", "Publication Track P2")
    add_cover(doc); add_contents(doc); add_manuscript_body(doc)
    doc.core_properties.title = "守恒型动态神经–SPH耦合的验证优先资格：中文稿v0.2"
    doc.core_properties.subject = "Publication Track P2 literature-positioned manuscript"
    doc.core_properties.author = "SPH-PIO-PoC Publication Workflow"
    doc.core_properties.keywords = "SPH, verification-first, literature verification, multistep gradient, topology event"
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    return {"output": str(OUTPUT), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


def build_references() -> dict:
    rows = list(csv.DictReader(CSV.open(encoding="utf-8-sig")))
    core = [r for r in rows if r["core_reference"].lower() == "true"]
    doc = Document(REFERENCE)
    clear_body(doc); repair_styles(doc)
    configure_header_footer(doc, "SPH-PIO-PoC  |  Verified Core References", "Publication Track P2")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    run = p.add_run("核验核心文献表"); set_font(run, 22, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Verified core bibliography · cutoff 2026-08-05 · n=40"); set_font(run, 10.5, color=MUTED)
    doc.add_paragraph("纳入标准：Crossref与出版社/官方会议页面双源核验；无独立DOI的会议论文使用官方会议页及作者预印本。序号与v0.2正文一致。")
    for r in core:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.28); p.paragraph_format.space_after = Pt(5)
        authors = r["authors"] or "NOT_REPORTED"
        venue = r["venue"] or "NOT_REPORTED"
        locator = f"https://doi.org/{r['doi']}" if r["doi"] else (r["publisher_url"] or r["secondary_url"])
        text = f"[{r['citation_id']}] {authors} ({r['year']}). {r['title']}. {venue}. {locator}"
        add_inline(p, text, 9.5)
    doc.core_properties.title = "Publication Track P2 verified core references"
    doc.save(REF_OUTPUT)
    return {"output": str(REF_OUTPUT), "references": len(core), "paragraphs": len(doc.paragraphs)}


def main() -> None:
    print(json.dumps({"manuscript": build_manuscript(), "references": build_references()}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
