#!/usr/bin/env python3
"""Build the Chinese Stage 03 Research Record from frozen Stage 03D-S evidence."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[3]
STAGE = REPO / "stage_03_Dynamic_SPH_Transformer_Hybrid"
ROOT = STAGE / "08_route_closure"
OUT = STAGE / "documents/Stage_03_Research_Record.docx"
ASSETS = ROOT / "figure_plan/record_assets"
TOC_PAGE_MAP = ROOT / "manifests/stage03ds_toc_pages.json"
SKILL = Path("/Users/xiejinbo/.codex/plugins/cache/openai-primary-runtime/documents/26.802.11031/skills/documents")
sys.path.insert(0, str(SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ledger = json.loads((ROOT / "status_ledger/stage03ds_status_ledger.json").read_text())
matrix = json.loads((ROOT / "evidence_matrix/stage03ds_dynamic_evidence_matrix.json").read_text())
gradient = json.loads((ROOT / "gradient_boundary/stage03ds_gradient_boundary.json").read_text())
topology = json.loads((ROOT / "topology_boundary/stage03ds_topology_component_boundary.json").read_text())
claims = json.loads((ROOT / "claim_boundary/stage03ds_claim_boundary.json").read_text())
assessment = json.loads((ROOT / "manuscript_assessment/stage03ds_manuscript_readiness.json").read_text())
figure_plan = json.loads((ROOT / "figure_plan/stage03ds_figure_and_table_plan.json").read_text())
future = json.loads((ROOT / "future_hypotheses/stage03ds_future_hypotheses.json").read_text())
freeze = json.loads((STAGE / "10_manifests/stage03ds_input_freeze_manifest.json").read_text())
toc_pages = json.loads(TOC_PAGE_MAP.read_text()) if TOC_PAGE_MAP.is_file() else {}


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183B56"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "2F6B4F"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Heiti SC"


def set_font(run, size=None, bold=None, color=None, italic=None, name=BODY_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT if name == BODY_FONT else name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_text(cell, text, *, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(str(text)), size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, weights, *, font_size=8.5, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
        shade(table.rows[0].cells[i], LIGHT_BLUE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            short = isinstance(value, (int, float, bool)) or len(str(value)) < 15
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER if short else WD_ALIGN_PARAGRAPH.LEFT, size=font_size)
            if zebra and row_index % 2:
                shade(cells[i], "FAFBFC")
    apply_table_geometry(table, column_widths_from_weights(weights, 9360), table_width_dxa=9360, indent_dxa=120, cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120})
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True, color=NAVY)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))
    return p


def add_equation(doc, expression, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(f"{expression}    ({label})"), size=11.2, italic=True, name="Cambria Math")


def add_callout(doc, label, text, color=NAVY):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(label + "  "), size=10, bold=True, color=color)
    set_font(p.add_run(text), size=10)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=9, color=MID_GRAY, italic=True)


def add_picture(doc, path, alt_text, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)


def add_page_field(paragraph):
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r._r.extend([begin, instr, separate, text, end])
    set_font(r, size=9, color=MID_GRAY)


bookmark_id = 10


def add_bookmark(paragraph, name):
    global bookmark_id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    bookmark_id += 1


def add_heading(doc, text, level, bookmark=None):
    p = doc.add_heading(text, level=level)
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_internal_link(paragraph, text, anchor):
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    link.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), NAVY); rpr.append(color)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    link.append(r)
    paragraph._p.append(link)


def add_external_link(paragraph, text, url):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLUE); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    link.append(r)
    paragraph._p.append(link)


def add_toc_entry(doc, number, title, bookmark):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8640")
    tabs.append(tab)
    p._p.get_or_add_pPr().append(tabs)
    add_internal_link(p, f"{number}. {title}", bookmark)
    set_font(p.add_run("\t" + str(toc_pages.get(bookmark, "—"))), size=8.8, color=NAVY)


def font_path():
    candidates = ["/System/Library/Fonts/PingFang.ttc", "/System/Library/Fonts/STHeiti Light.ttc", "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"]
    return next(path for path in candidates if Path(path).is_file())


FONT_PATH = font_path()


def fonts(size):
    return ImageFont.truetype(FONT_PATH, size, index=0)


def centered(draw, box, text, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=font, fill=fill)


def pipeline_figure(path):
    img = Image.new("RGB", (1800, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((65, 35), "Stage 03 verification-first 路线与关闭边界", font=fonts(46), fill="#183B56")
    items = [
        ("03A", "Specification", "PASS", "#E8EEF5", "#2E74B5"),
        ("03B", "References", "PASS", "#E8EEF5", "#2E74B5"),
        ("03C", "Implementation", "VERIFIED", "#EAF5EF", "#2F6B4F"),
        ("03D", "Multistep AD/FD", "NOT QUALIFIED", "#FFF4F4", "#9B1C1C"),
        ("03D-R", "Attribution", "MIXED / UNRESOLVED", "#FFF8E8", "#7A5A00"),
        ("03D-S", "Route closure", "PAUSED", "#F2F4F7", "#183B56"),
    ]
    x0, y, w, h, gap = 55, 210, 245, 280, 43
    for i, (stage, name, verdict, fill, outline) in enumerate(items):
        x = x0 + i * (w + gap)
        d.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline=outline, width=4)
        centered(d, (x, y + 18, x + w, y + 82), stage, fonts(32), outline)
        centered(d, (x + 10, y + 95, x + w - 10, y + 155), name, fonts(23), "#263238")
        for j, line in enumerate(verdict.split(" / ")):
            centered(d, (x + 8, y + 175 + j * 38, x + w - 8, y + 213 + j * 38), line, fonts(20), outline)
        if i < len(items) - 1:
            d.line((x + w + 4, y + h / 2, x + w + gap - 10, y + h / 2), fill="#6B7280", width=5)
            d.polygon([(x + w + gap - 10, y + h / 2), (x + w + gap - 28, y + h / 2 - 10), (x + w + gap - 28, y + h / 2 + 10)], fill="#6B7280")
    d.text((65, 625), "独立 topology component: QUALIFIED   |   Stage 03E authorization: false   |   training / rollout: not executed", font=fonts(25), fill="#9B1C1C")
    img.save(path)


def architecture_figure(path):
    img = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((65, 35), "D0-D3 architecture 与共同守恒边界", font=fonts(46), fill="#183B56")
    rows = [
        ("D0", "Baseline WCSPH", "无 correction / 无 memory"),
        ("D1", "Instantaneous pair MLP", "current token only"),
        ("D2", "Causal recurrent pair PIO", "shared recurrent state"),
        ("D3", "Causal reciprocal Transformer", "H=4 temporal attention"),
    ]
    for i, (arm, name, temporal) in enumerate(rows):
        y = 165 + i * 155
        d.rounded_rectangle((70, y, 390, y + 110), radius=15, fill="#E8EEF5", outline="#2E74B5", width=3)
        centered(d, (70, y, 390, y + 110), f"{arm}\n{name}", fonts(25), "#183B56")
        d.rounded_rectangle((530, y, 900, y + 110), radius=15, fill="#F2F4F7", outline="#6B7280", width=3)
        centered(d, (530, y, 900, y + 110), temporal, fonts(24), "#263238")
        d.rounded_rectangle((1060, y, 1695, y + 110), radius=15, fill="#EAF5EF", outline="#2F6B4F", width=3)
        centered(d, (1060, y, 1695, y + 110), "reciprocal antisymmetric pair-force head\n(D1-D3; D0 is zero correction)", fonts(22), "#2F6B4F")
        d.line((390, y + 55, 530, y + 55), fill="#6B7280", width=5)
        d.line((900, y + 55, 1060, y + 55), fill="#6B7280", width=5)
    d.text((70, 815), "共同边界：相同合法输入与结构门；没有训练结果，也没有 D3 优越性前提。", font=fonts(25), fill="#9B1C1C")
    img.save(path)


def outcomes_figure(path):
    img = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((65, 35), "360-probe 多步 AD/FD 完整结果与失败归因", font=fonts(46), fill="#183B56")
    x0, y0, total_w, h = 130, 170, 1540, 115
    pass_w = total_w * 216 / 360
    d.rounded_rectangle((x0, y0, x0 + total_w, y0 + h), radius=12, fill="#FFF4F4", outline="#9B1C1C", width=3)
    d.rounded_rectangle((x0, y0, x0 + pass_w, y0 + h), radius=12, fill="#EAF5EF", outline="#2F6B4F", width=3)
    centered(d, (x0, y0, x0 + pass_w, y0 + h), "PASS 216 (60%)", fonts(30), "#2F6B4F")
    centered(d, (x0 + pass_w, y0, x0 + total_w, y0 + h), "FAIL 144 (40%)", fonts(30), "#9B1C1C")
    reasons = list(gradient["failure_reason_counts"].items())
    max_count = max(v for _, v in reasons)
    for i, (name, count) in enumerate(reasons):
        y = 365 + i * 75
        d.text((100, y), name, font=fonts(20), fill="#263238")
        bw = 850 * count / max_count
        d.rounded_rectangle((760, y, 760 + bw, y + 38), radius=8, fill="#D98C8C", outline="#9B1C1C", width=2)
        d.text((780 + bw, y + 2), str(count), font=fonts(21), fill="#9B1C1C")
    d.text((100, 905), "完整资格化失败；D-R 归因为 mixed or unresolved，并未覆盖 Stage 03D。", font=fonts(25), fill="#9B1C1C")
    img.save(path)


def claim_figure(path):
    img = Image.new("RGB", (1800, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((65, 35), "Stage 03 证据—主张边界", font=fonts(46), fill="#183B56")
    columns = [
        ("SUPPORTED", "#EAF5EF", "#2F6B4F", ["RK2 implementation verified", "bitwise zero correction", "structural conservation", "TE1 topology qualified", "complete gradient qualification not achieved"]),
        ("CONDITIONAL", "#FFF8E8", "#7A5A00", ["D3 backend sensitivity", "history attenuation", "FD conditioning contribution", "no systematic vanish/explode detected"]),
        ("UNSUPPORTED", "#FFF4F4", "#9B1C1C", ["Transformer trainable", "rollout improves SPH", "D3 outperforms D1/D2", "cutoff membership differentiable", "long-time stability"]),
    ]
    for i, (head, fill, outline, items) in enumerate(columns):
        x = 65 + i * 575
        d.rounded_rectangle((x, 150, x + 520, 770), radius=20, fill=fill, outline=outline, width=4)
        centered(d, (x, 175, x + 520, 235), head, fonts(30), outline)
        y = 285
        for item in items:
            d.ellipse((x + 32, y + 8, x + 46, y + 22), fill=outline)
            d.text((x + 65, y), item, font=fonts(21), fill="#263238")
            y += 88
    img.save(path)


ASSETS.mkdir(parents=True, exist_ok=True)
asset_paths = [
    ASSETS / "record_figure_01_pipeline.png",
    ASSETS / "record_figure_02_architecture.png",
    ASSETS / "record_figure_03_adfd_outcomes.png",
    ASSETS / "record_figure_04_claim_boundary.png",
]
pipeline_figure(asset_paths[0])
architecture_figure(asset_paths[1])
outcomes_figure(asset_paths[2])
claim_figure(asset_paths[3])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = BODY_FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for name, size, color, before, after in (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, DARK_BLUE, 10, 5)):
    style = styles[name]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[list_name].font.name = BODY_FONT
    styles[list_name]._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    styles[list_name]._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    styles[list_name]._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    styles[list_name].font.size = Pt(11)
    styles[list_name].paragraph_format.space_after = Pt(4)
    styles[list_name].paragraph_format.line_spacing = 1.25
styles["Caption"].font.name = BODY_FONT
styles["Caption"]._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
styles["Caption"]._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
styles["Caption"].font.size = Pt(9)
styles["Caption"].font.italic = True
styles["Caption"].font.color.rgb = RGBColor.from_string(MID_GRAY)

settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
set_font(hp.add_run("SPH-PIO-PoC  |  Stage 03 Research Record"), size=9, color=MID_GRAY)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Stage 03 研究记录  |  "), size=9, color=MID_GRAY)
add_page_field(fp)

# editorial_cover pattern
doc.add_paragraph().paragraph_format.space_after = Pt(82)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("SPH-PIO-PoC"), size=12, bold=True, color=GOLD)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(10)
set_font(p.add_run("Stage 03 研究记录"), size=30, bold=True, color=NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("保守动态 neural-SPH：实现验证、拓扑事件与多步梯度边界"), size=15, color=DARK_BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(40)
set_font(p.add_run("Dynamic Route Closure, Evidence Synthesis and Publication Boundary"), size=10.5, italic=True, color=MID_GRAY)
add_callout(doc, "终端边界", "Stage 03C implementation verified；Stage 03D multistep gradient NOT QUALIFIED；Stage 03D-R mixed or unresolved；topology component qualified；Stage 03E authorization=false。", RED)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(28)
set_font(p.add_run("证据冻结日期：2026-08-05  |  非计算性闭包  |  训练与 rollout 均未执行"), size=10, color=MID_GRAY)
add_bookmark(doc.paragraphs[0], "Top")

doc.add_page_break()
toc_heading = add_heading(doc, "目录", 1, "TOC")
toc_entries = [
    ("1", "Stage 03 新假设", "sec01"), ("2", "Stage 02 边界", "sec02"),
    ("3", "动态控制方程", "sec03"), ("4", "D0-D3 架构", "sec04"),
    ("5", "RK2 / history / graph semantics", "sec05"), ("6", "动态参考体系", "sec06"),
    ("7", "Stage 03B trajectory qualification", "sec07"), ("8", "Stage 03C implementation", "sec08"),
    ("9", "Zero correction", "sec09"), ("10", "Conservation / equivariance", "sec10"),
    ("11", "Checkpoint / resume", "sec11"), ("12", "One-step autograd", "sec12"),
    ("13", "Stage 03D multistep AD/FD", "sec13"), ("14", "TE1 topology qualification", "sec14"),
    ("15", "Stage 03D-R failure attribution", "sec15"), ("16", "Supported / unsupported claims", "sec16"),
    ("17", "Publication boundary", "sec17"), ("18", "Future new hypotheses", "sec18"),
    ("19", "Artifact / hash index", "sec19"),
]
for number, title, anchor in toc_entries:
    add_toc_entry(doc, number, title, anchor)
add_body(doc, "目录条目包含内部链接；页码由最终渲染后的静态索引生成。")

doc.add_page_break()
add_heading(doc, "摘要", 1, "abstract")
add_body(doc, "Stage 03 以“动态、因果、守恒的 particle interaction correction 可否在严格 V&V 合同下实现并进入训练”为新假设。路线获得了强实现证据：独立 RK2、bitwise zero correction、结构守恒/等变、checkpoint/resume 与 one-step autograd 全部通过；同时，冻结的 360-probe 多步 AD/FD 合同仅有216个 stable windows，144个 failure，故完整梯度资格化失败。Stage 03D-R 进一步表明 failure 来源混合且部分未决，不能覆盖 Stage 03D 的 NOT_QUALIFIED。TE1 拓扑事件分量独立通过，但 cutoff edge existence 本身仍为离散、不可微事件。")
add_callout(doc, "闭包结论", "动态学习路线暂停于多步梯度边界。训练、solver-in-the-loop、rollout 性能和长期稳定性均未授权或未测试；不得写为“训练失败”或“Transformer 求解器全部失败”。", RED)
add_caption(doc, "图 1  Stage 03 verification-first 路线与关闭边界")
add_picture(doc, asset_paths[0], "Stage 03A through Stage 03D-S verification-first workflow with preserved failure and no Stage 03E authorization")

add_heading(doc, "1. Stage 03 新假设", 1, "sec01")
add_body(doc, "Stage 03 的新假设不是“Transformer 必然改进 SPH”，而是：在 baseline WCSPH 完整保留的条件下，新增的因果、局部、互易 pair-force correction 能否首先满足可审计的控制方程、时间积分、历史因果性、图拓扑、守恒/等变、checkpoint 与 differentiability 合同，然后才具备进入训练和 rollout 的资格。")
add_body(doc, "该假设被拆成可证伪的层级：specification → reference → implementation → multistep gradient/topology → dataset/training。Stage 03D 的失败触发了路线暂停，而不是把前层级已获得的实现证据一并否定。")

add_heading(doc, "2. Stage 02 边界", 1, "sec02")
add_body(doc, "Stage 02 静态学习路线已以 STAGE02_ROUTE_CLOSED_PUBLICATION_BOUNDARY_COMPLETE 关闭。Stage 02M 与 Stage 02M-Q 两个冻结静态拟合协议均未资格化；Stage 02 的 consumed test、selected checkpoints 和 K2 结果只能作为历史静态诊断。Stage 03 不复用其权重，不把静态 v0.2 延伸为 v0.3，也不以 rollout 隐藏静态 failure。")
add_callout(doc, "历史边界", "Stage 01 仍为 V2_QUALIFICATION_FAIL；Stage 01H 仍为 FINITE_RESOLUTION_DOMINANT；viscosity operator form 仍为 NOT_CONFIRMED。Stage 03 未恢复或改写这些状态。", RED)

add_heading(doc, "3. 动态控制方程", 1, "sec03")
add_body(doc, "冻结的 Stage 03 v0.1 半离散系统保留 baseline WCSPH 的连续性、压力、EOS、图和时间步选择；网络只提供加性 acceleration correction。")
add_equation(doc, "dx_i/dt = v_i", "1")
add_equation(doc, "dρ_i/dt = C_SPH,i(S)", "2")
add_equation(doc, "dv_i/dt = a_SPH,i(S) + a_θ,i(S_history, H_history, G_history)", "3")
add_equation(doc, "p_i = c_s²(ρ_i − ρ₀)", "4")
add_equation(doc, "a_θ,i = (1/m_i) Σ_{j:{i,j}∈G} f_θ,ij", "5")
add_equation(doc, "f_θ,ij = F⁰_ij[α_ij r̂_ij + β_ij t_ij],  α_ij=α_ji, β_ij=β_ji", "6")
add_body(doc, "由于 r̂_ji=−r̂_ij、t_ji=−t_ij 且标量系数对称，pair force 满足 f_θ,ji=−f_θ,ij。该构造建立结构守恒，不建立 correction 的准确性或必要性。")

add_heading(doc, "4. D0-D3 架构", 1, "sec04")
add_caption(doc, "图 2  D0-D3 architecture 与共同守恒边界")
add_picture(doc, asset_paths[1], "D0 baseline, D1 instantaneous MLP, D2 recurrent PIO, and D3 causal temporal reciprocal Transformer architecture comparison")
add_table(doc, ["Arm", "角色", "时间机制", "Pair-force head", "证据边界"], [
    ["D0", "Baseline WCSPH", "none", "correction disabled", "zero/cost baseline"],
    ["D1", "Instantaneous conservative pair MLP", "current token", "shared antisymmetric basis", "memory necessity control"],
    ["D2", "Causal recurrent pair PIO", "recurrent state", "shared antisymmetric basis", "non-attention temporal control"],
    ["D3", "Causal temporal reciprocal Transformer PIO", "H=4 causal Transformer", "shared antisymmetric basis", "candidate, not presumed winner"],
], [0.55, 1.85, 1.25, 1.45, 1.4], font_size=7.9)

add_heading(doc, "5. RK2 / history / graph semantics", 1, "sec05")
add_heading(doc, "5.1 Explicit midpoint/RK2", 2)
add_equation(doc, "k₁ = F(Sⁿ, historyⁿ)", "7")
add_equation(doc, "Sⁿ⁺¹ᐟ² = Sⁿ + (Δt/2)k₁", "8")
add_equation(doc, "k₂ = F(Sⁿ⁺¹ᐟ², historyⁿ + ephemeral token)", "9")
add_equation(doc, "Sⁿ⁺¹ = Sⁿ + Δt k₂", "10")
add_body(doc, "start 与 midpoint 是两个独立 RHS evaluation；每次都重算 EOS 并从相应 state 重建 reciprocal graph。固定 whole-step topology 被禁止。")
add_heading(doc, "5.2 Transactional history", 2)
add_body(doc, "start/midpoint 只读 committed snapshot。midpoint token 为 ephemeral，不得 append、evict 或覆盖 accepted history。只有 Sⁿ⁺¹ 通过 finite/safety checks 后，才在物理时刻 tₙ₊₁ 构造一个 accepted token 并原子 commit；失败步同时回滚 state 与 cache。")
add_heading(doc, "5.3 Topology boundary", 2)
add_body(doc, "Graph construction 由 baseline solver 控制。网络不得预测、剪枝或改变 edge membership。edge birth/death 由 cutoff crossing 离散触发；事件两侧可讨论固定拓扑路径的梯度，事件本身不被声明可微。")

add_heading(doc, "6. 动态参考体系", 1, "sec06")
add_table(doc, ["Class", "作用", "本阶段状态", "禁止越界"], [
    ["D-R1", "analytic/MMS verification", "two families PASS", "MMS ≠ physical validation"],
    ["D-R2", "same-semidiscrete high-accuracy time reference", "six cases PASS", "time reference ≠ spatial truth"],
    ["D-R3", "source-free analytic/independent validation", "oblique shear PASS", "never training/normalization/threshold source"],
    ["D-R4", "external V&V-qualified reference", "NOT_AVAILABLE", "higher-resolution SPH alone ≠ D-R4"],
], [0.65, 2.15, 1.45, 2.25], font_size=8.2)

add_heading(doc, "7. Stage 03B trajectory qualification", 1, "sec07")
add_body(doc, "Stage 03B 在 deterministic CPU float64 下完成参考资格化。D-R1 Lagrangian compression 与 coupled deformation 通过解析闭包；D-R2 对六个同半离散算子 case 完成 DOP853 time-reference sensitivity；D-R3 oblique shear A/B 通过 source-free exact audit。最终形成18条 canonical trajectories。")
add_table(doc, ["Evidence", "Result", "Boundary"], [
    ["D-R1", "2 families; 6 exact trajectories; PASS", "verification/MMS only"],
    ["D-R2", "6/6 time-reference cases PASS", "not spatial truth"],
    ["D-R3 oblique shear", "2 families; 6 exact trajectories; PASS", "independent validation only"],
    ["Acoustic", "linear-regime conditional", "not unrestricted exact D-R3"],
    ["Periodic vortex", "rejected as exact source-free", "possible separate MMS role only"],
    ["Topology events", "0 in Stage 03B trajectories", "TE1 handled later"],
], [1.35, 2.35, 2.8], font_size=8.2)
add_body(doc, "17 frames constitute one trajectory rather than IID samples。Stage 03B 未建立 dataset split、normalization、optimizer 或 neural rollout；其53.49 s wall time、348,471,296-byte peak RSS 和4302 RHS/rebuilds仅为资格化资源记录。")

add_heading(doc, "8. Stage 03C implementation", 1, "sec08")
add_body(doc, "Stage 03C 以 Stage 03B 为唯一授权，完成 D0-D3 frozen implementation contract。其 final status 为 DYNAMIC_RK2_HYBRID_IMPLEMENTATION_VERIFIED。该状态表示实现及指定 smoke/structural gates 通过，不表示 multistep gradient、training 或 performance 通过。")
add_table(doc, ["Gate", "Frozen result", "Interpretation"], [
    ["Independent RK2", "48/48 PASS", "state/time/graph semantics"],
    ["Zero correction", "288/288 bitwise PASS", "D1-D3 zero head equals D0"],
    ["History", "PASS", "single accepted commit; no midpoint contamination"],
    ["Structure", "PASS", "antisymmetry/conservation/equivariance"],
    ["Checkpoint/resume", "6/6 PASS", "state/graph/history/RNG identity"],
    ["One-step autograd", "6/6 PASS", "plumbing only"],
    ["Resources/safety", "PASS", "not performance benchmark"],
], [1.55, 1.8, 3.15], font_size=8.3)

add_heading(doc, "9. Zero correction", 1, "sec09")
add_body(doc, "Zero correction 使用两种冻结 mode：Mode A 完全 bypass neural correction；Mode B 执行网络但将 correction head 置零。D1-D3、多个参考 case 与 horizons 上的 accepted state、density、pressure、velocity、wrapped/unwrapped position、graph hashes、source count 和 step/time 均与 D0 bitwise 一致，合计288/288。")
add_callout(doc, "允许主张", "zero-correction equivalence is bitwise established。不得外推为 nonzero learned correction accurate 或 dynamic solver improved。", GREEN)

add_heading(doc, "10. Conservation / equivariance", 1, "sec10")
add_body(doc, "Stage 03C structural smoke 对 pair exchange、force antisymmetry、normalized correction-force residual、deterministic repeat、edge reorder、particle permutation、translation、periodic representative shift、Galilean boost、SO(2) rotation 与 reflection 进行检查并通过。Stage 03D 的多步审计进一步记录540/540 per-stage conservation PASS。")
add_body(doc, "这些结果说明 reciprocal pair-force 结构在多个 RK stages 中保持预期不变量；它们不是 long-time stability、accuracy 或 learned improvement 的替代证据。")

add_heading(doc, "11. Checkpoint / resume", 1, "sec11")
add_body(doc, "D0、D1 zero-head、D2 zero-head、D3 zero-head、D2 fixed-seed 与 D3 fixed-seed 六个 configuration 在 save/load 后均复现 physical state、accepted time/step、parameter hash、history order/hash、hidden state、RNG identity 与 graph hash sequence。Stage 03C 没有训练 checkpoint；这里验证的是 deterministic continuation plumbing。")

add_heading(doc, "12. One-step autograd", 1, "sec12")
add_body(doc, "Stage 03C 对 D1-D3 × 两个 family 共6个 one-step runs 检查 expected parameter、initial state 与 history-related gradients；全部 finite、nonzero 且 deterministic repeat PASS。没有 edge-index gradient、optimizer object 或 parameter update，也没有 finite difference。")
add_callout(doc, "层级边界", "one-step autograd verified ≠ complete multistep differentiability qualified。后者由 Stage 03D 的冻结 AD/FD 合同单独判定。", GOLD)

doc.add_page_break()
add_heading(doc, "13. Stage 03D multistep AD/FD", 1, "sec13")
add_caption(doc, "图 3  360-probe 多步 AD/FD 完整结果与失败归因")
add_picture(doc, asset_paths[2], "Complete 360-probe multistep AD/FD outcome with 216 pass and 144 failure rows plus failure-reason counts")
add_body(doc, "冻结 matrix 覆盖 D1-D3、四个 case roles、三个 seeds、horizons 1/2/4/8 与参数/初值/history probes，共360 rows、2880 AD/FD comparisons。216 rows 获得 stable adjacent epsilon window，144 rows failure；没有 topology-changing epsilon 被排除。")
add_table(doc, ["Arm", "PASS", "FAIL", "Total"], [["D1", 65, 31, 96], ["D2", 75, 45, 120], ["D3", 76, 68, 144], ["All", 216, 144, 360]], [1.8, 1.5, 1.5, 1.7], font_size=9)
add_table(doc, ["Primary reason", "Count"], [[k, v] for k, v in gradient["failure_reason_counts"].items()], [5.3, 1.2], font_size=8.5)
add_callout(doc, "Stage 03D verdict", "DYNAMIC_MULTISTEP_ADFD_AND_TOPOLOGY_NOT_QUALIFIED。不得只展示216个PASS，不得隐藏144个failure。", RED)

add_heading(doc, "14. TE1 topology qualification", 1, "sec14")
add_body(doc, "TE1 对 cutoff crossing 建立独立 component contract：记录一个 edge birth 与一个 death，6/6 stage replay PASS，12/12 event-side AD/FD PASS，并确认 force jump finite/bounded 与 empty graph deterministic semantics。")
add_table(doc, ["Topology gate", "Result", "Boundary"], [
    ["Birth/death", "1 / 1", "deterministic event semantics"],
    ["Replay", "6/6 PASS", "frozen TE1 only"],
    ["Fixed-side gradients", "12/12 PASS", "not through membership change"],
    ["Force jump", "finite/bounded", "no global smoothness claim"],
    ["Empty graph", "deterministic", "no correction edges"],
], [2.0, 1.45, 3.05], font_size=8.5)
add_callout(doc, "Component verdict", "TOPOLOGY_EVENT_COMPONENT_QUALIFIED。该 PASS 不把 Stage 03D 总体状态改成 PASS，也不证明 arbitrary topology families 或 cutoff existence differentiable。", GREEN)

add_heading(doc, "15. Stage 03D-R failure attribution", 1, "sec15")
add_body(doc, "Stage 03D-R 不修改 Stage 03D contract、epsilon、probe、backend 或 architecture。它首先重建 complete 360-row matrix，再对冻结 selected rows 执行 same-math reverse/JVP、extended FD、objective decomposition、history path 与 horizon scaling。")
add_table(doc, ["Diagnostic", "Frozen result", "Interpretation"], [
    ["Same-math reverse/JVP", "60/60 PASS", "AD implementation consistency"],
    ["Historical backend vs math JVP", "48/60 match", "12 selected D3 mismatches; backend sensitivity"],
    ["Extended FD", "30/60 stable; 2640 paths", "conditioning contributes to some failures"],
    ["History traces", "1 conditioning-limited; 5 below FD resolution", "strong rollout attenuation"],
    ["Horizon scaling", "90 bounded/nonmonotone", "no systematic vanish/explode detected"],
    ["Unresolved", "19/144", "no complete single-root-cause attribution"],
], [2.1, 2.05, 2.35], font_size=8.1)
add_body(doc, "Temporal-module-only paths generally exhibited stable reverse/JVP/FD behavior, while full-rollout history influence was attenuated to ratios approximately 1.54×10⁻⁴–4.91×10⁻³。No detach/cache/object mismatch was identified。该证据支持 mixed contributors，而不支持“全部 failure 都是 FD artifact”或“D3 intrinsically non-differentiable”。")
add_callout(doc, "D-R verdict", "DYNAMIC_GRADIENT_FAILURE_MIXED_OR_UNRESOLVED。D-R 不覆盖、不修复 Stage 03D failure。", RED)

doc.add_page_break()
add_heading(doc, "16. Supported / unsupported claims", 1, "sec16")
add_caption(doc, "图 4  Stage 03 证据—主张边界")
add_picture(doc, asset_paths[3], "Supported, conditional, and unsupported claim map for Stage 03")
add_heading(doc, "16.1 Supported claims", 2)
for row in claims["supported_claims"]:
    add_body(doc, "允许：" + row["allowed_wording"] + " 禁止：" + row["prohibited_wording"])
add_heading(doc, "16.2 Conditional claims", 2)
for row in claims["conditional_claims"]:
    add_body(doc, "条件性允许：" + row["allowed_wording"] + " 禁止：" + row["prohibited_wording"])
add_heading(doc, "16.3 Unsupported claims", 2)
unsupported_rows = [[row["claim"], row["allowed_wording"], row["prohibited_wording"]] for row in claims["unsupported_claims"]]
add_table(doc, ["Unsupported claim", "Allowed wording", "Prohibited wording"], unsupported_rows, [1.75, 2.8, 1.95], font_size=7.8)
add_callout(doc, "语义守卫", "未执行动态训练只能写 NOT AUTHORIZED / NOT EXECUTED，不能写“动态训练失败”；Stage 03D NOT_QUALIFIED 不能写“Transformer 求解器全部失败”。", RED)

add_heading(doc, "17. Publication boundary", 1, "sec17")
add_body(doc, "当前不能形成以完整动态 solver 性能为中心的论文，因为没有训练、受控/自主 rollout、独立性能比较或 long-time stability evidence。当前最可辩护的主线是 verification-first conservative dynamic neural-SPH coupling：把分层合同、bitwise zero correction、结构守恒、TE1 event boundary 和完整负梯度证据作为方法贡献。")
add_table(doc, ["Paper", "方向", "Readiness", "判断"], [[r["paper"], r["direction"], r["readiness"], r["assessment"]] for r in assessment["papers"]], [0.55, 2.1, 1.4, 2.45], font_size=8.0)
add_body(doc, "CMAME 的官方 scope 包括 meshless discretization、fluid mechanics 和 physically based machine learning，因此主题层面相符；但该刊强调具有显著发展的原创 computational methods。当前证据不足以支撑 full-solver claim，Paper B 仍需更强的通用方法深度与独立验证，Paper C 则需证明诊断框架超出单个 PoC。")
p = doc.add_paragraph(); set_font(p.add_run("官方 scope："), bold=True, color=NAVY); add_external_link(p, "Computer Methods in Applied Mechanics and Engineering — Aims & Scope", assessment["cmame_scope_source"])
add_heading(doc, "17.1 三项核心缺失证据", 2)
for item in assessment["missing_core_evidence"]:
    add_bullet(doc, item)
add_heading(doc, "17.2 主文、补充材料与内部证据", 2)
add_body(doc, "主文：" + "；".join(assessment["main_text_results"]) + "。")
add_body(doc, "补充材料：" + "；".join(assessment["supplementary_audits"]) + "。")
add_body(doc, "内部证据：" + "；".join(assessment["internal_only"]) + "。")

add_heading(doc, "18. Future new hypotheses", 1, "sec18")
add_body(doc, "以下仅为新假设设计，Stage 03D-S 不执行。所有分支都必须作为全新 Stage 04 建立新的 hypothesis、implementation/topology/AD-FD contract 与 success gates；不得称为 Stage 03E 直接继续。")
add_table(doc, ["ID", "Future hypothesis", "Required stage", "Executed"], [[f"H{r['id']}", r["hypothesis"], r["required_stage"], "No"] for r in future["hypotheses"]], [0.5, 3.75, 1.65, 0.6], font_size=8.2)

add_heading(doc, "19. Artifact / hash index", 1, "sec19")
add_body(doc, f"Stage 03D-S input freeze 在综合前记录 {freeze['historical_file_count']} 个历史文件。完整逐文件 SHA-256 位于 10_manifests/stage03ds_input_freeze_manifest.json；本节列出关键索引。")
key_paths = [
    "stage_03_Dynamic_SPH_Transformer_Hybrid/10_manifests/stage03a_final_manifest.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/10_manifests/stage03b_final_manifest.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/10_manifests/stage03c_final_manifest.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/10_manifests/stage03d_final_manifest.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/10_manifests/stage03dr_final_manifest.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03d/results/fixed_topology_adfd_results.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03dr/failure_matrix/stage03d_complete_360_row_matrix.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03dr/ad_crosscheck/reverse_vs_jvp.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03dr/fd_conditioning/extended_fd_results.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03dr/history_path/reference_prehistory_trace.json",
    "stage_03_Dynamic_SPH_Transformer_Hybrid/05_dynamic_solver_implementation/stage03dr/topology_preservation/topology_component_status.json",
]
frozen_lookup = {row["path"]: row for row in freeze["historical_files"]}
index_rows = []
for path in key_paths:
    row = frozen_lookup[path]
    index_rows.append([Path(path).name, path, row["sha256"][:30] + "…"])
add_table(doc, ["Artifact", "Path", "SHA-256 prefix"], index_rows[:6], [1.5, 4.0, 1.0], font_size=7.3)
doc.add_page_break()
add_heading(doc, "19.1 Artifact / hash index（续）", 2)
add_table(doc, ["Artifact", "Path", "SHA-256 prefix"], index_rows[6:], [1.5, 4.0, 1.0], font_size=7.3)

add_heading(doc, "19.2 Stage 03 状态账本", 2)
add_table(doc, ["Stage", "Unique status", "Exec", "Steps", "Runs", "Performance"], [[r["stage"], r["status"], r["execution_count"], r["optimizer_steps"], r["training_runs"], r["performance_evaluations"]] for r in ledger["rows"]], [0.75, 3.1, 0.6, 0.65, 0.65, 0.75], font_size=7.8)

doc.add_page_break()
add_heading(doc, "附录 A：动态证据矩阵", 1, "appendixA")
matrix_rows = [[r["id"], r["category"].replace("_", " "), r["item"], r["status"], r["boundary"]] for r in matrix["rows"]]
add_table(doc, ["ID", "Category", "Item", "Status", "Boundary"], matrix_rows[:15], [0.45, 1.35, 1.55, 1.05, 2.1], font_size=7.6)
doc.add_page_break()
add_heading(doc, "附录 A.1：动态证据矩阵（续）", 2)
add_table(doc, ["ID", "Category", "Item", "Status", "Boundary"], matrix_rows[15:], [0.45, 1.35, 1.55, 1.05, 2.1], font_size=7.6)

doc.add_page_break()
add_heading(doc, "附录 B：图表发表计划", 1, "appendixB")
add_table(doc, ["Fig.", "Title", "Form", "Integrity rule"], [[r["figure"], r["title"], r["form"], r["integrity_rule"]] for r in figure_plan["figures"]], [0.45, 2.2, 1.3, 2.55], font_size=7.8)
add_body(doc, "计划表格：" + "；".join(row["title"] for row in figure_plan["tables"]) + "。不得生成虚构训练或性能图。")

add_heading(doc, "记录闭包声明", 1, "closure")
add_callout(doc, "Stage 03D-S", "本研究记录不授权新 AD/FD contract、epsilon、probe、backend、architecture、dataset、training protocol、optimizer、training、rollout 或 solver-in-the-loop。最终闭包状态须以渲染审计与 final manifest 为准。", NAVY)
add_bookmark(doc.paragraphs[-1], "Bottom")

doc.core_properties.title = "Stage 03 Research Record"
doc.core_properties.subject = "Dynamic route closure, evidence synthesis and publication boundary"
doc.core_properties.author = "SPH-PIO-PoC Research Record"
doc.core_properties.keywords = "SPH, conservative neural solver, RK2, multistep gradient, topology event, verification"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(json.dumps({"output": str(OUT), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "figures": len(doc.inline_shapes), "toc_pages_loaded": bool(toc_pages)}, ensure_ascii=False))
