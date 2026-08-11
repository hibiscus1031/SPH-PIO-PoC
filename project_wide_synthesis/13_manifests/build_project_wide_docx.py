#!/usr/bin/env python3
"""Create the Chinese project-wide research synthesis DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "project_wide_synthesis"
DEST = OUT / "documents" / "SPH_PIO_PoC_Project_Wide_Research_Synthesis.docx"
FONT = "Source Han Sans CN"
NAVY = "17324D"; BLUE = "2E74B5"; DARK = "1F4D78"; MUTED = "5E6B76"; PALE = "F4F6F9"


def set_font(run, name=FONT, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:cantSplit"); el.set(qn("w:val"), "true"); trPr.append(el)


def set_repeat_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None: layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    ind = tblPr.find(qn("w:tblInd"))
    if ind is None: ind = OxmlElement("w:tblInd"); tblPr.append(ind)
    ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            w = widths[i]
            cell.width = Inches(w / 1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for mname in ("top","bottom","start","end"):
                tcMar = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
                if tcMar is None: tcMar=OxmlElement("w:tcMar"); cell._tc.get_or_add_tcPr().append(tcMar)
                m=tcMar.find(qn(f"w:{mname}"))
                if m is None: m=OxmlElement(f"w:{mname}"); tcMar.append(m)
                m.set(qn("w:w"), "80" if mname in ("top","bottom") else "120"); m.set(qn("w:type"), "dxa")


def bookmark(paragraph, name, bid):
    start=OxmlElement("w:bookmarkStart"); start.set(qn("w:id"),str(bid)); start.set(qn("w:name"),name)
    end=OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"),str(bid))
    paragraph._p.insert(0,start); paragraph._p.append(end)


def internal_link(paragraph, anchor, text):
    h=OxmlElement("w:hyperlink"); h.set(qn("w:anchor"),anchor)
    r=OxmlElement("w:r"); rPr=OxmlElement("w:rPr")
    color=OxmlElement("w:color"); color.set(qn("w:val"),BLUE); rPr.append(color)
    u=OxmlElement("w:u"); u.set(qn("w:val"),"single"); rPr.append(u)
    fonts=OxmlElement("w:rFonts"); fonts.set(qn("w:eastAsia"),FONT); fonts.set(qn("w:ascii"),FONT); fonts.set(qn("w:hAnsi"),FONT); rPr.append(fonts)
    r.append(rPr); t=OxmlElement("w:t"); t.text=text; r.append(t); h.append(r); paragraph._p.append(h)


def add_field(paragraph, instruction):
    run=paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=instruction
    sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate")
    txt=OxmlElement("w:t"); txt.text="1"; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    run._r.extend([begin,instr,sep,txt,end]); set_font(run,size=9,color=MUTED)


def add_table(doc, caption, headers, rows, widths, anchor, bid):
    cp=doc.add_paragraph(style="Caption"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=cp.add_run(caption); set_font(rr,size=9,bold=True,color=DARK); bookmark(cp,anchor,bid)
    cp.paragraph_format.space_before=Pt(4); cp.paragraph_format.space_after=Pt(4)
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; set_repeat_table_geometry(t,widths)
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=str(h); set_cell_shading(t.rows[0].cells[i],"E8EEF5")
        for r in t.rows[0].cells[i].paragraphs[0].runs: set_font(r,size=8,bold=True,color=NAVY)
    mark_header(t.rows[0])
    prevent_row_split(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            for p in cells[i].paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.0
                for r in p.runs:set_font(r,size=7.5,color="222222")
        prevent_row_split(t.rows[-1])
    set_repeat_table_geometry(t,widths)
    return t


doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(0.492)

styles=doc.styles
normal=styles["Normal"]; normal.font.name=FONT; normal._element.rPr.rFonts.set(qn("w:eastAsia"),FONT); normal.font.size=Pt(11)
normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; normal.paragraph_format.space_before=Pt(0); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333
for name,size,color,before,after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    st=styles[name]; st.font.name=FONT; st._element.rPr.rFonts.set(qn("w:eastAsia"),FONT); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
cap=styles["Caption"]; cap.font.name=FONT; cap._element.rPr.rFonts.set(qn("w:eastAsia"),FONT); cap.font.size=Pt(9); cap.font.color.rgb=RGBColor.from_string(DARK)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
hr=header.add_run("SPH-PIO-PoC｜Project-Wide Evidence Synthesis"); set_font(hr,size=8.5,color=MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=footer.add_run("Cross-Stage Synthesis S1  ·  "); set_font(fr,size=8.5,color=MUTED); add_field(footer,"PAGE")

# Editorial cover.
sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(80)
k=doc.add_paragraph(); k.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(k.add_run("PROJECT-WIDE RESEARCH SYNTHESIS"),size=11,bold=True,color=BLUE)
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_after=Pt(10)
set_font(t.add_run("SPH-PIO-PoC\n全项目研究综合与发表决策档案"),size=28,bold=True,color=NAVY)
st=doc.add_paragraph(); st.alignment=WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_after=Pt(42)
set_font(st.add_run("Cross-Stage Synthesis S1 — Evidence Audit, Failure Analysis, Innovation Register and Publication Decision Dossier"),size=12,color=DARK)
meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(meta.add_run("只读、非计算性审计｜扫描截止 2026-08-05\nGit HEAD ff86f5e0b99966ad6fa5896fe3d9a0c3f001cd57\n文档预设：narrative_proposal；命名字体覆盖：Source Han Sans CN"),size=9.5,color=MUTED)
doc.add_page_break()

md=(OUT/"12_reports/project_wide_research_synthesis.md").read_text(encoding="utf-8")
parts=re.split(r"^## ",md,flags=re.M)[1:]
parsed=[]
for part in parts:
    lines=part.splitlines(); title=lines[0].strip(); body="\n".join(lines[1:]).strip()
    if title.startswith("附录"): continue
    parsed.append((title,body))

toc=doc.add_paragraph(style="Heading 1"); toc.add_run("目录"); bookmark(toc,"TOC",1)
for i,(title,_) in enumerate(parsed,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15); p.paragraph_format.space_after=Pt(3)
    internal_link(p,f"sec{i:03d}",title)
doc.add_page_break()

bid=10
for idx,(title,body) in enumerate(parsed,1):
    h=doc.add_paragraph(style="Heading 1"); h.add_run(title); bookmark(h,f"sec{idx:03d}",bid); bid+=1
    # Strip markdown adornment and split into readable paragraphs.
    clean=body.replace("**","").replace("`","")
    paras=[x.strip() for x in re.split(r"\n\s*\n",clean) if x.strip()]
    for txt in paras:
        if txt.startswith("|") or txt.startswith("#"): continue
        p=doc.add_paragraph(); set_font(p.add_run(txt),size=10.5,color="222222")
    if title.startswith("3."):
        fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(fp.add_run("资格链：L0 → L1 → … → L10；任一局部 PASS 不覆盖上游或总体 FAIL。"),size=10,italic=True,color=DARK)
    if title.startswith("9."):
        p=doc.add_paragraph(); p.add_run("跨阶段static PIO闭环见"); internal_link(p,"fig1","图 1"); p.add_run("。")
    if title.startswith("11."):
        pic=doc.add_picture(str(ROOT/"stage_02_Particle_Interaction_Operator/08_route_closure/figure_plan/record_assets/record_figure_01_pipeline.png"),width=Inches(6.45))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        pic._inline.docPr.set("descr","Stage 02从reference、target、dataset、architecture到两版static fitting失败与route closure的资格流程")
        cp=doc.add_paragraph(style="Caption"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(cp.add_run("图 1｜Stage 02 verification-first PIO资格流程与终止边界"),size=9,bold=True,color=DARK); bookmark(cp,"fig1",bid); bid+=1
    if title.startswith("16."):
        eq=doc.add_paragraph(); eq.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(eq.add_run("stable probe ⇔ ∃ adjacent ε window satisfying magnitude, direction and consistency gates"),name="Menlo",size=9.5,color=DARK)
        pic=doc.add_picture(str(ROOT/"stage_03_Dynamic_SPH_Transformer_Hybrid/08_route_closure/figure_plan/record_assets/record_figure_03_adfd_outcomes.png"),width=Inches(6.45))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        pic._inline.docPr.set("descr","360个多步AD/FD probe中216通过、144失败，并按方向、结构零、非单调、舍入、截断、非光滑和未解析分类")
        cp=doc.add_paragraph(style="Caption"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(cp.add_run("图 2｜360-probe多步AD/FD完整结果与失败归因"),size=9,bold=True,color=DARK); bookmark(cp,"fig2",bid); bid+=1
    if title.startswith("28."):
        claims=json.loads((OUT/"07_claim_boundary/project_wide_claim_boundary.json").read_text())["claims"]
        p=doc.add_paragraph(); p.add_run("最终允许/禁止措辞见"); internal_link(p,"tbl2","表 2"); p.add_run("。")
        add_table(doc,"表 2｜全项目主张边界",["ID","分类","允许措辞","禁止措辞"],[(x["id"],x["classification"],x["allowed_wording"],x["prohibited_wording"]) for x in claims],[700,1500,3600,3560],"tbl2",bid); bid+=1

# Appendices with compact machine-ledger tables.
timeline=json.loads((OUT/"02_stage_timeline/complete_stage_timeline.json").read_text())["rows"]
h=doc.add_paragraph(style="Heading 1"); h.add_run("附录 A｜阶段状态与证据索引"); bookmark(h,"secA",bid); bid+=1
summary=[]
for r in timeline:
    if r["stage_id"] in {"Stage 00","Stage 01B","Stage 01D2","Stage 01G execution","Stage 01H","Stage 02K","Stage 02M","Stage 02M-Q","Stage 02M-S","Stage 03A","Stage 03B","Stage 03C","Stage 03D","Stage 03D-R","Stage 03D-S","Publication P1","Publication P2"}:
        summary.append((r["stage_id"],r["exact_final_status"],r["blocker"],r["authorized_input"]))
add_table(doc,"表 1｜关键里程碑状态账本",["阶段","exact status","主要阻断","机器/冻结证据"],summary,[1150,2550,2600,3060],"tbl1",bid); bid+=1

fail=json.loads((OUT/"04_failure_register/complete_failure_register.json").read_text())["events"]
h=doc.add_paragraph(style="Heading 1"); h.add_run("附录 B｜A–R失败类别总览"); bookmark(h,"secB",bid); bid+=1
add_table(doc,"表 3｜失败类别、历史状态与方法学后果",["ID/类别","阶段/状态","直接原因","方法学后果"],[(x["id"]+" "+x["category"],x["stage"]+" / "+x["exact_status"],x["direct_cause"],x["downstream_effect"]) for x in fail],[1900,2400,2600,2460],"tbl3",bid); bid+=1

h=doc.add_paragraph(style="Heading 1"); h.add_run("附录 C｜文档导航与可访问性说明"); bookmark(h,"secC",bid); bid+=1
p=doc.add_paragraph(); set_font(p.add_run("目录条目使用内部书签链接；页脚包含动态PAGE字段；图1和图2均写入替代文本；表1–3首行标记为重复表头。正文引用图表使用内部链接。字体采用Source Han Sans CN命名覆盖以稳定中文渲染。"),size=10.5)
p=doc.add_paragraph(); internal_link(p,"TOC","返回目录")

DEST.parent.mkdir(parents=True,exist_ok=True)
doc.save(DEST)
print(json.dumps({"output":str(DEST),"sections":len(parsed),"tables":3,"figures":2},ensure_ascii=False))
